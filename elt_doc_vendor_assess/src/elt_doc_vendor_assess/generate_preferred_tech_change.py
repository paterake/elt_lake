"""Generate a Preferred Technologies Change .docx document for ACB approval.

Usage:
    uv run --package elt-doc-vendor-assess python elt_doc_vendor_assess/src/elt_doc_vendor_assess/generate_preferred_tech_change.py <json_file> <output_path>

Input JSON structure:
{
    "technology_name": "Hyve Managed SFTP Hosting",
    "technology_type": "SaaS",   // one of: "SaaS", "Application", "Library"
    "document_history": [
        {"version": "1.0", "date": "2026-02-10", "author": "Rakesh Patel", "summary": "Initial version"}
    ],
    "reference_documents": [
        {"document": "SAD Document", "link": "..."}
    ],
    "sections": {
        "functionality": {
            "key_features": "...",
            "benefits_vs_competitors": "...",
            "advantages_over_mainstream": "..."
        },
        "fit": {
            "interoperability": "...",
            "customization": "..."
        },
        "security_saas": {
            "commercial_or_open_source": "...",
            "community_support": "...",
            "iso_27001": "...",
            "soc2": "...",
            "eu_hosting": "...",
            "patching": "...",
            "vulnerability_scans": "...",
            "end_user_auth": "...",
            "admin_access": "...",
            "vendor_data_access": "...",
            "encryption": "...",
            "breach_history": "...",
            "high_availability": "...",
            "backup": "..."
        },
        "security_application": {
            "commercial_or_open_source": "...",
            "community_support": "...",
            "iso_27001": "...",
            "soc2": "...",
            "eu_hosting": "...",
            "patching": "...",
            "vulnerability_scans": "...",
            "end_user_auth": "...",
            "admin_access": "...",
            "vendor_data_access": "..."
        },
        "security_library": {
            "commercial_or_open_source": "...",
            "community_support": "..."
        },
        "versioning": {
            "version_control": "..."
        },
        "hosting": {
            "infrastructure": "..."
        },
        "maintenance": {
            "roadmap": "...",
            "release_cycle": "...",
            "upgrade_process": "..."
        },
        "knowledge": {
            "training": "..."
        },
        "implementation": {
            "implementation_plan": "...",
            "decommissioning_plan": "...",
            "migration_activities": "..."
        },
        "costs": {
            "licensing": "...",
            "upgrade_support": "..."
        }
    },
    "appendix": "..."
}
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


# ── Section definitions ──────────────────────────────────────────────────────

SECTIONS = [
    ("1. Functionality", "functionality", [
        ("1.1", "Key Features", "key_features",
         "Describe the key features and functionality of the technology."),
        ("1.2", "Benefits vs. Competitors", "benefits_vs_competitors",
         "What are the benefits over competing products?"),
        ("1.3", "Advantages over Existing Mainstream", "advantages_over_mainstream",
         "What advantages does this technology offer over the existing Mainstream technology?"),
    ]),
    ("2. Fit", "fit", [
        ("2.1", "Interoperability", "interoperability",
         "How does the technology interoperate with existing FA systems?"),
        ("2.2", "Customization", "customization",
         "How are customizations handled?"),
    ]),
    ("3. Security — SaaS Product", "security_saas", [
        ("3.1", "Commercial or Open Source", "commercial_or_open_source",
         "Is the product commercial or open source?"),
        ("3.2", "Community / Vendor Support", "community_support",
         "What support community or vendor support is available?"),
        ("3.3", "ISO 27001 Certification", "iso_27001",
         "Does the vendor hold ISO 27001 certification?"),
        ("3.4", "SOC 2 Type II", "soc2",
         "Does the vendor hold SOC 2 Type II certification?"),
        ("3.5", "EU/UK Hosting", "eu_hosting",
         "Is data hosted in the EU or UK?"),
        ("3.6", "Patching", "patching",
         "What is the patching strategy and cadence?"),
        ("3.7", "Vulnerability Scanning", "vulnerability_scans",
         "How are vulnerability scans performed?"),
        ("3.8", "End User Authentication", "end_user_auth",
         "What authentication mechanisms are available for end users?"),
        ("3.9", "Admin Access", "admin_access",
         "How is administrative access managed?"),
        ("3.10", "Vendor Data Access", "vendor_data_access",
         "What access does the vendor have to customer data?"),
        ("3.11", "Encryption", "encryption",
         "Describe encryption at rest and in transit."),
        ("3.12", "Breach History", "breach_history",
         "Any security breaches in the last 5 years?"),
        ("3.13", "High Availability", "high_availability",
         "Describe the high availability architecture."),
        ("3.14", "Backup & Recovery", "backup",
         "Describe backup procedures and recovery processes."),
    ]),
    ("4. Security — Application", "security_application", [
        ("4.1", "Commercial or Open Source", "commercial_or_open_source",
         "Is the application commercial or open source?"),
        ("4.2", "Community / Vendor Support", "community_support",
         "What support community or vendor support is available?"),
        ("4.3", "ISO 27001 Certification", "iso_27001",
         "Does the vendor hold ISO 27001 certification?"),
        ("4.4", "SOC 2 Type II", "soc2",
         "Does the vendor hold SOC 2 Type II certification?"),
        ("4.5", "EU/UK Hosting", "eu_hosting",
         "Is data hosted or processed in the EU/UK?"),
        ("4.6", "Patching", "patching",
         "What is the patching strategy?"),
        ("4.7", "Vulnerability Scanning", "vulnerability_scans",
         "How are vulnerability scans performed?"),
        ("4.8", "End User Authentication", "end_user_auth",
         "What authentication mechanisms are available?"),
        ("4.9", "Admin Access", "admin_access",
         "How is admin access managed?"),
        ("4.10", "Vendor Data Access", "vendor_data_access",
         "Does the vendor have access to data processed by the application?"),
    ]),
    ("5. Security — Library or Component", "security_library", [
        ("5.1", "Commercial or Open Source", "commercial_or_open_source",
         "Is the library commercial or open source?"),
        ("5.2", "Community Support", "community_support",
         "What is the size and activity of the community?"),
    ]),
    ("6. Versioning", "versioning", [
        ("6.1", "Version Control", "version_control",
         "How are artefacts and configurations version-controlled?"),
    ]),
    ("7. Hosting", "hosting", [
        ("7.1", "Infrastructure Requirements", "infrastructure",
         "What infrastructure is required?"),
    ]),
    ("8. Maintenance", "maintenance", [
        ("8.1", "Roadmap", "roadmap",
         "Describe the product roadmap."),
        ("8.2", "Release Cycle", "release_cycle",
         "What is the release cycle?"),
        ("8.3", "Upgrade Process", "upgrade_process",
         "How are upgrades applied?"),
    ]),
    ("9. Knowledge", "knowledge", [
        ("9.1", "Training Requirements", "training",
         "What training is required for staff?"),
    ]),
    ("10. Implementation", "implementation", [
        ("10.1", "Implementation Plan", "implementation_plan",
         "Provide a high-level implementation plan."),
        ("10.2", "Decommissioning Plan", "decommissioning_plan",
         "If replacing an existing technology, describe the decommissioning plan."),
        ("10.3", "Migration Activities", "migration_activities",
         "What migration activities are required?"),
    ]),
    ("11. Costs", "costs", [
        ("11.1", "Licensing Costs", "licensing",
         "Detail the licensing costs."),
        ("11.2", "Upgrade & Support Costs", "upgrade_support",
         "Detail ongoing upgrade and support costs."),
    ]),
]

# Map technology_type to the security section that should be populated
SECURITY_SECTIONS = {
    "SaaS": "security_saas",
    "Application": "security_application",
    "Library": "security_library",
}


# ── Styling helpers ──────────────────────────────────────────────────────────

def _set_cell_text(cell, text, bold=False, font_size=9):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"


def _shade_cell(cell, color_hex):
    """Apply background shading to a cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _style_header_row(row, bg_color="1F4E79"):
    """Style a table header row."""
    for cell in row.cells:
        _shade_cell(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def _add_heading(doc, text, level):
    """Add a heading with consistent font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


# ── Main document builder ────────────────────────────────────────────────────

def generate_preferred_tech_change(data: dict, output_path: str):
    """Generate the Preferred Technologies Change .docx document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    tech_type = data.get("technology_type", "SaaS")
    active_security_section = SECURITY_SECTIONS.get(tech_type, "security_saas")

    # ── Title ────────────────────────────────────────────────────────────
    _add_heading(doc, f"Preferred Technologies Change — {data['technology_name']}", level=0)

    # ── Document History Table ───────────────────────────────────────────
    _add_heading(doc, "Document History", level=2)
    history = data.get("document_history", [])
    if history:
        cols = ["Version No", "Revision Date", "Author", "Summary of Changes"]
        table = doc.add_table(rows=1 + len(history), cols=len(cols))
        table.style = "Table Grid"
        for i, col in enumerate(cols):
            _set_cell_text(table.rows[0].cells[i], col, bold=True)
        _style_header_row(table.rows[0])
        for r, h in enumerate(history, start=1):
            _set_cell_text(table.rows[r].cells[0], h.get("version", ""))
            _set_cell_text(table.rows[r].cells[1], h.get("date", ""))
            _set_cell_text(table.rows[r].cells[2], h.get("author", ""))
            _set_cell_text(table.rows[r].cells[3], h.get("summary", ""))
    doc.add_paragraph()

    # ── Reference Documents Table ────────────────────────────────────────
    refs = data.get("reference_documents", [])
    if refs:
        _add_heading(doc, "Reference Documents", level=2)
        table = doc.add_table(rows=1 + len(refs), cols=2)
        table.style = "Table Grid"
        _set_cell_text(table.rows[0].cells[0], "Document", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Link", bold=True)
        _style_header_row(table.rows[0])
        for r, ref in enumerate(refs, start=1):
            _set_cell_text(table.rows[r].cells[0], ref.get("document", ""))
            _set_cell_text(table.rows[r].cells[1], ref.get("link", ""))
        doc.add_paragraph()

    # ── Sections ─────────────────────────────────────────────────────────
    sections_data = data.get("sections", {})

    for section_heading, section_key, sub_questions in SECTIONS:
        _add_heading(doc, section_heading, level=1)

        # Handle security sections — only populate the relevant one
        is_security = section_key.startswith("security_")
        if is_security and section_key != active_security_section:
            doc.add_paragraph(
                f"N/A — this is a {tech_type} product. "
                f"See section for Security — {tech_type} above."
            )
            continue

        section_content = sections_data.get(section_key, {})

        if isinstance(section_content, str):
            doc.add_paragraph(section_content)
            continue

        for num, sub_heading, field_key, question_text in sub_questions:
            _add_heading(doc, f"{num} {sub_heading}", level=2)
            answer = section_content.get(field_key, "")
            if answer:
                for para in str(answer).split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph(f"[{question_text}]")

    # ── Appendix ─────────────────────────────────────────────────────────
    _add_heading(doc, "Appendix", level=1)
    appendix = data.get("appendix", "")
    if appendix:
        for para in str(appendix).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── CLI entry point ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: generate_preferred_tech_change.py <json_file> <output_path>")
        sys.exit(1)

    json_path = Path(sys.argv[1])
    output = sys.argv[2]

    with open(json_path) as f:
        data = json.load(f)

    generate_preferred_tech_change(data, output)
