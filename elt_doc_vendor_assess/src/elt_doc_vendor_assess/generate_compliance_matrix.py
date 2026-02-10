"""Generate a Vendor Compliance Matrix .docx document.

Usage:
    uv run --package elt-doc-vendor-assess python elt_doc_vendor_assess/src/elt_doc_vendor_assess/generate_compliance_matrix.py <json_file> <output_path>

Input JSON structure:
{
    "title": "SaaS SFTP Vendor Compliance Assessment",
    "category": "SaaS SFTP",
    "executive_summary": "...",
    "strategic_context": {
        "intro": "...",
        "technologies": [{"technology": "...", "classification": "...", "guidance": "..."}],
        "scope_definition": "...",
        "justification": "..."
    },
    "mandatory_requirements": {
        "intro": "...",
        "requirements": ["ISO 27001 OR SOC 2 Type II", "UK/EU data residency", ...]
    },
    "vendors": [
        {
            "name": "Hyve Managed Hosting",
            "tool_name": "Cerberus SFTP",
            "tool_type": "Managed SFTP",
            "deployment": "Private Cloud (UK)",
            "maturity": "Established",
            "iso_27001": true, "iso_27002": false, "iso_27017": true, "iso_27018": false,
            "iso_9001": true, "iso_22301": false, "iso_27701": false,
            "soc2": true, "soc1": false, "hipaa": false, "gdpr": true,
            "pci_dss": false, "fedramp": false, "ccpa": false, "itar": false,
            "cloud_provider": "Private Infrastructure",
            "uk_eu_hosting": true,
            "annual_cost": "£4,252",
            "setup_cost": "£750",
            "first_year_cost": "£5,002",
            "rbac": true, "abac": false, "ssh_key_mgmt": true,
            "auto_key_rotation": false, "mfa": true, "sso": false,
            "encryption_at_rest": true, "tenancy_model": "Dedicated",
            "ip_whitelisting": true, "data_archiving": false, "auto_purging": true,
            "audit_logging": true, "log_retention": "90 days", "dlp": false,
            "malware_scan": false, "edr": false,
            "gdpr_art_5_6": true, "gdpr_art_15_17": true, "gdpr_art_25": true,
            "gdpr_art_30": true, "gdpr_art_32": true, "gdpr_art_33": true,
            "panorays_compatible": true,
            "encryption_transit": true, "credential_mgmt": "SSH Keys",
            "tier": 1,
            "tier_justification": "Meets all 4 mandatory requirements..."
        }
    ],
    "recommended_vendor": {
        "name": "Hyve Managed Hosting",
        "key_strengths": ["...", "..."],
        "contracted_config": "...",
        "sla_commitments": "..."
    },
    "tier_assessments": {
        "tier_1": [{"name": "...", "summary": "...", "label": "PRIMARY RECOMMENDATION"}],
        "tier_2": [{"name": "...", "summary": "..."}],
        "tier_3": [{"name": "...", "summary": "..."}]
    },
    "recommendations_summary": "...",
    "conclusion": "...",
    "certification_methodology": "...",
    "domain_context": {                          // OPTIONAL
        "heading": "Workday Integration Context",
        "subsections": [
            {"heading": "Rate Types Required", "content": "..."},
            {"heading": "Key Integration Notes", "bullets": ["...", "..."]}
        ]
    },
    "incumbent_assessment": {                    // OPTIONAL
        "heading": "Microsoft FX Rate API Assessment",
        "intro": "...",
        "finding_heading": "Finding: No Suitable Microsoft FX Rate API Exists",
        "finding_text": "...",
        "table": [
            {"service": "...", "description": "...", "limitation": "...", "compatible": "No"}
        ]
    },
    "include_tables": {                          // OPTIONAL — all default to true
        "soc_regulatory": false,
        "infra_cost": true,
        "advanced_security": false,
        "gdpr_articles": false
    },
    "verification_sources": [{"vendor": "...", "url": "..."}]
}
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ── Styling helpers ──────────────────────────────────────────────────────────

def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"


def _bool_to_symbol(val):
    """Convert boolean/string to checkmark/X/dash."""
    if isinstance(val, bool):
        return "✓" if val else "✗"
    if isinstance(val, str):
        return val  # already formatted (e.g. "Via Azure", "Unknown")
    return "—"


def _shade_cell(cell, color_hex):
    """Apply background shading to a cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _style_header_row(row, bg_color="1F4E79"):
    """Style a table header row with dark background and white text."""
    for cell in row.cells:
        _shade_cell(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def _add_heading(doc, text, level):
    """Add a heading with consistent font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


# ── Table builders ───────────────────────────────────────────────────────────

def _build_overview_table(doc, vendors, extra_columns=None):
    """Provider Overview & Capabilities. Extra columns are domain-specific."""
    _add_heading(doc, "Provider Overview & Capabilities", level=2)
    base_cols = ["Vendor", "Tool Name", "Tool Type", "Deployment", "Maturity"]
    extra = extra_columns or []  # list of {"header": "...", "field": "..."}
    col_headers = base_cols + [ec["header"] for ec in extra]
    table = doc.add_table(rows=1 + len(vendors), cols=len(col_headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(col_headers):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(table.rows[r].cells[1], v.get("tool_name", "—"), align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(table.rows[r].cells[2], v.get("tool_type", "—"))
        _set_cell_text(table.rows[r].cells[3], v.get("deployment", "—"))
        _set_cell_text(table.rows[r].cells[4], v.get("maturity", "—"))
        for c, ec in enumerate(extra, start=5):
            val = v.get(ec["field"], "—")
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(val) if isinstance(val, bool) else str(val or "—"))

    doc.add_paragraph()


def _build_iso_table(doc, vendors):
    """ISO/SOC 2 Compliance Matrix with Compliance Route."""
    _add_heading(doc, "ISO/SOC 2 Compliance Matrix", level=2)
    # Build field list dynamically — only include ISO columns that have data
    all_iso_fields = [
        ("iso_27001", "ISO 27001"), ("iso_27002", "ISO 27002"),
        ("iso_27017", "ISO 27017"), ("iso_27018", "ISO 27018"),
        ("iso_9001", "ISO 9001"), ("iso_22301", "ISO 22301"),
        ("iso_27701", "ISO 27701"),
    ]
    # Always include 27001; include others only if at least one vendor has data
    iso_fields = [f for f in all_iso_fields if f[0] == "iso_27001" or any(v.get(f[0]) is not None for v in vendors)]

    # Always include SOC 2, GDPR, and Compliance Route
    cert_fields = [("soc2", "SOC 2\nType II"), ("gdpr", "GDPR")]
    route_col = True  # always include Compliance Route

    cols = ["Vendor"] + [label for _, label in iso_fields] + [label for _, label in cert_fields] + (["Compliance\nRoute"] if route_col else [])
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        c = 0
        _set_cell_text(table.rows[r].cells[c], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        c += 1
        for field, _ in iso_fields:
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(v.get(field)))
            c += 1
        for field, _ in cert_fields:
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(v.get(field)))
            c += 1
        if route_col:
            _set_cell_text(table.rows[r].cells[c], v.get("compliance_route", "—"), align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()


def _build_soc_regulatory_table(doc, vendors):
    """Table 3: SOC 2, HIPAA, GDPR & Regional Compliance."""
    _add_heading(doc, "SOC 2, HIPAA, GDPR & Regional Compliance", level=2)
    fields = [
        ("soc2", "SOC 2\nType II"), ("hipaa", "HIPAA"), ("gdpr", "GDPR"),
        ("pci_dss", "PCI DSS"), ("fedramp", "FedRAMP"), ("ccpa", "CCPA"),
        ("itar", "ITAR"),
    ]
    cols = ["Vendor"] + [label for _, label in fields]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        for c, (field, _) in enumerate(fields, start=1):
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(v.get(field)))

    doc.add_paragraph()


def _build_infra_cost_table(doc, vendors):
    """Table 4: Infrastructure & Cost Comparison."""
    _add_heading(doc, "Infrastructure & Cost Comparison", level=2)
    cols = ["Vendor", "Cloud / Infrastructure", "UK/EU Hosting", "Est. Annual Cost", "Setup Cost", "First Year Total"]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(table.rows[r].cells[1], v.get("cloud_provider", "—"))
        _set_cell_text(table.rows[r].cells[2], _bool_to_symbol(v.get("uk_eu_hosting")))
        _set_cell_text(table.rows[r].cells[3], v.get("annual_cost", "—"))
        _set_cell_text(table.rows[r].cells[4], v.get("setup_cost", "—"))
        _set_cell_text(table.rows[r].cells[5], v.get("first_year_cost", "—"))

    doc.add_paragraph()


def _build_security_features_table(doc, vendors, custom_fields=None):
    """Security Features Matrix with flexible columns."""
    _add_heading(doc, "Security Features Matrix", level=2)
    # Default fields for infrastructure products; override with custom_fields if provided
    fields = custom_fields or [
        ("rbac", "RBAC"), ("abac", "ABAC"), ("ssh_key_mgmt", "SSH Key\nMgmt"),
        ("auto_key_rotation", "Auto Key\nRotation"), ("mfa", "MFA/2FA"),
        ("sso", "SSO"), ("encryption_at_rest", "Encryption\nat Rest"),
        ("tenancy_model", "Tenancy\nModel"),
    ]
    cols = ["Vendor"] + [label for _, label in fields]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        for c, (field, _) in enumerate(fields, start=1):
            val = v.get(field)
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(val) if isinstance(val, bool) else str(val or "—"))

    doc.add_paragraph()


def _build_advanced_security_table(doc, vendors):
    """Table 6: Advanced Security Features."""
    _add_heading(doc, "Advanced Security Features", level=2)
    fields = [
        ("ip_whitelisting", "IP\nWhitelisting"), ("data_archiving", "Data\nArchiving"),
        ("auto_purging", "Auto\nPurging"), ("audit_logging", "Audit\nLogging"),
        ("log_retention", "Log\nRetention"), ("dlp", "DLP"),
        ("malware_scan", "Malware\nScan"), ("edr", "EDR"),
    ]
    cols = ["Vendor"] + [label for _, label in fields]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        for c, (field, _) in enumerate(fields, start=1):
            val = v.get(field)
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(val) if isinstance(val, bool) else str(val or "—"))

    doc.add_paragraph()


def _build_gdpr_table(doc, vendors):
    """Table 7: GDPR Article Compliance."""
    _add_heading(doc, "GDPR Article Compliance", level=2)
    fields = [
        ("gdpr_art_5_6", "Art. 5,6\n(Lawfulness)"),
        ("gdpr_art_15_17", "Art. 15,17\n(Rights)"),
        ("gdpr_art_25", "Art. 25\n(By Design)"),
        ("gdpr_art_30", "Art. 30\n(Records)"),
        ("gdpr_art_32", "Art. 32\n(Security)"),
        ("gdpr_art_33", "Art. 33\n(Breach)"),
    ]
    cols = ["Vendor"] + [label for _, label in fields]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    for r, v in enumerate(vendors, start=1):
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        for c, (field, _) in enumerate(fields, start=1):
            _set_cell_text(table.rows[r].cells[c], _bool_to_symbol(v.get(field)))

    doc.add_paragraph()


def _build_compliance_summary_table(doc, vendors):
    """Table 8: Compliance Assessment Summary."""
    _add_heading(doc, "Compliance Assessment Summary", level=2)
    cols = ["Vendor", "ISO/SOC 2", "UK/EU", "Cloud/Infra", "Panorays", "Status"]
    table = doc.add_table(rows=1 + len(vendors), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])

    tier_labels = {1: "COMPLIANT", 2: "REVIEW", 3: "NON-COMPLIANT"}
    tier_colors = {1: "C6EFCE", 2: "FFEB9C", 3: "FFC7CE"}

    for r, v in enumerate(vendors, start=1):
        has_iso_soc = v.get("iso_27001") or v.get("soc2")
        _set_cell_text(table.rows[r].cells[0], v["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(table.rows[r].cells[1], _bool_to_symbol(has_iso_soc))
        _set_cell_text(table.rows[r].cells[2], _bool_to_symbol(v.get("uk_eu_hosting")))
        _set_cell_text(table.rows[r].cells[3], _bool_to_symbol(bool(v.get("cloud_provider"))))
        _set_cell_text(table.rows[r].cells[4], _bool_to_symbol(v.get("panorays_compatible")))
        tier = v.get("tier", 3)
        status_cell = table.rows[r].cells[5]
        _set_cell_text(status_cell, tier_labels.get(tier, "—"), bold=True)
        _shade_cell(status_cell, tier_colors.get(tier, "FFFFFF"))

    doc.add_paragraph()


# ── Conditional section builders ──────────────────────────────────────────────

def _build_domain_context(doc, ctx):
    """Build an Integration / Domain Context section."""
    _add_heading(doc, ctx.get("heading", "Integration Context"), level=1)
    for sub in ctx.get("subsections", []):
        _add_heading(doc, sub["heading"], level=2)
        if sub.get("content"):
            doc.add_paragraph(sub["content"])
        for bullet in sub.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")


def _build_incumbent_assessment(doc, inc):
    """Build an Incumbent / Platform Vendor Assessment section with table."""
    _add_heading(doc, inc.get("heading", "Incumbent Vendor Assessment"), level=1)
    if inc.get("intro"):
        doc.add_paragraph(inc["intro"])
    if inc.get("finding_heading"):
        _add_heading(doc, inc["finding_heading"], level=2)
    if inc.get("finding_text"):
        doc.add_paragraph(inc["finding_text"])
    rows = inc.get("table", [])
    if rows:
        cols = ["Service / Product", "Description", "Limitation", "Compatible?"]
        table = doc.add_table(rows=1 + len(rows), cols=len(cols))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, col in enumerate(cols):
            _set_cell_text(table.rows[0].cells[i], col, bold=True)
        _style_header_row(table.rows[0])
        for r, row in enumerate(rows, start=1):
            _set_cell_text(table.rows[r].cells[0], row.get("service", ""), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(table.rows[r].cells[1], row.get("description", ""), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(table.rows[r].cells[2], row.get("limitation", ""), align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(table.rows[r].cells[3], row.get("compatible", "No"))
        doc.add_paragraph()


def _build_tier3_summary_table(doc, tier3_entries):
    """Build a compact Tier 3 Non-Compliant summary table."""
    if len(tier3_entries) < 3:
        return False  # use prose for 1-2 entries
    cols = ["Provider", "Assessment Notes", "Status"]
    table = doc.add_table(rows=1 + len(tier3_entries), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(cols):
        _set_cell_text(table.rows[0].cells[i], col, bold=True)
    _style_header_row(table.rows[0])
    for r, entry in enumerate(tier3_entries, start=1):
        _set_cell_text(table.rows[r].cells[0], entry["name"], align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(table.rows[r].cells[1], entry.get("summary", ""), align=WD_ALIGN_PARAGRAPH.LEFT)
        status_cell = table.rows[r].cells[2]
        _set_cell_text(status_cell, "NON-COMPLIANT", bold=True)
        _shade_cell(status_cell, "FFC7CE")
    doc.add_paragraph()
    return True


# ── Main document builder ────────────────────────────────────────────────────

def generate_compliance_matrix(data: dict, output_path: str):
    """Generate the Vendor Compliance Matrix .docx document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Determine which conditional tables to include
    incl = data.get("include_tables", {})
    include_soc_regulatory = incl.get("soc_regulatory", True)
    include_infra_cost = incl.get("infra_cost", True)
    include_advanced_security = incl.get("advanced_security", True)
    include_gdpr = incl.get("gdpr_articles", True)

    # ── Title ────────────────────────────────────────────────────────────
    _add_heading(doc, data["title"], level=0)

    # ── Executive Summary ────────────────────────────────────────────────
    _add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(data.get("executive_summary", ""))

    # ── Recommended Vendor ───────────────────────────────────────────────
    rec = data.get("recommended_vendor", {})
    if rec:
        _add_heading(doc, f"Recommended Vendor: {rec['name']}", level=2)
        if rec.get("key_strengths"):
            _add_heading(doc, "Key Strengths", level=3)
            for strength in rec["key_strengths"]:
                doc.add_paragraph(strength, style="List Bullet")
        if rec.get("contracted_config"):
            _add_heading(doc, "Contracted Deployment Configuration", level=3)
            doc.add_paragraph(rec["contracted_config"])
        if rec.get("sla_commitments"):
            _add_heading(doc, "Service Level Commitments", level=3)
            doc.add_paragraph(rec["sla_commitments"])

    # ── Domain Context (conditional) ─────────────────────────────────────
    if data.get("domain_context"):
        _build_domain_context(doc, data["domain_context"])

    # ── Incumbent Assessment (conditional) ───────────────────────────────
    if data.get("incumbent_assessment"):
        _build_incumbent_assessment(doc, data["incumbent_assessment"])

    # ── Strategic Alignment ──────────────────────────────────────────────
    sc = data.get("strategic_context")
    if sc:
        _add_heading(doc, "Strategic Alignment with FA Preferred Technologies", level=1)
        if sc.get("intro"):
            doc.add_paragraph(sc["intro"])

        if sc.get("technologies"):
            _add_heading(doc, "Relevant Technology Classifications", level=2)
            tech_table = doc.add_table(rows=1 + len(sc["technologies"]), cols=3)
            tech_table.style = "Table Grid"
            for i, col in enumerate(["Technology", "Classification", "Strategic Guidance"]):
                _set_cell_text(tech_table.rows[0].cells[i], col, bold=True)
            _style_header_row(tech_table.rows[0])
            for r, t in enumerate(sc["technologies"], start=1):
                _set_cell_text(tech_table.rows[r].cells[0], t["technology"], align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell_text(tech_table.rows[r].cells[1], t["classification"])
                _set_cell_text(tech_table.rows[r].cells[2], t["guidance"], align=WD_ALIGN_PARAGRAPH.LEFT)
            doc.add_paragraph()

        if sc.get("scope_definition"):
            _add_heading(doc, "Scope Definition", level=2)
            doc.add_paragraph(sc["scope_definition"])

        if sc.get("justification"):
            _add_heading(doc, "Justification", level=2)
            doc.add_paragraph(sc["justification"])

    # ── FA Security Mandatory Requirements ───────────────────────────────
    mreq = data.get("mandatory_requirements", {})
    _add_heading(doc, "FA Security Mandatory Requirements", level=1)
    if mreq.get("intro"):
        doc.add_paragraph(mreq["intro"])
    for req in mreq.get("requirements", []):
        doc.add_paragraph(req, style="List Number")

    # ── Vendor Compliance Matrix (tables) ────────────────────────────────
    _add_heading(doc, "Vendor Compliance Matrix", level=1)
    vendors = data.get("vendors", [])
    extra_cols = data.get("overview_extra_columns", [])
    _build_overview_table(doc, vendors, extra_columns=extra_cols or None)
    _build_iso_table(doc, vendors)
    if include_soc_regulatory:
        _build_soc_regulatory_table(doc, vendors)
    if include_infra_cost:
        _build_infra_cost_table(doc, vendors)
    security_fields = data.get("security_features_fields")
    _build_security_features_table(doc, vendors, custom_fields=security_fields)
    if include_advanced_security:
        _build_advanced_security_table(doc, vendors)
    if include_gdpr:
        _build_gdpr_table(doc, vendors)

    # ── Compliance Assessment Summary ────────────────────────────────────
    _add_heading(doc, "FA Security Compliance Assessment", level=1)
    _build_compliance_summary_table(doc, vendors)

    # ── Tier Assessments ─────────────────────────────────────────────────
    tiers = data.get("tier_assessments", {})
    if tiers.get("tier_1"):
        _add_heading(doc, "Tier 1 — FULLY COMPLIANT", level=2)
        for entry in tiers["tier_1"]:
            label = entry.get("label", entry["name"])
            _add_heading(doc, label, level=3)
            doc.add_paragraph(entry.get("summary", ""))

    if tiers.get("tier_2"):
        _add_heading(doc, "Tier 2 — REQUIRES REVIEW", level=2)
        for entry in tiers["tier_2"]:
            _add_heading(doc, entry["name"], level=3)
            doc.add_paragraph(entry.get("summary", ""))

    if tiers.get("tier_3"):
        _add_heading(doc, "Tier 3 — NON-COMPLIANT", level=2)
        # Use summary table for 3+ entries, prose for fewer
        if not _build_tier3_summary_table(doc, tiers["tier_3"]):
            for entry in tiers["tier_3"]:
                _add_heading(doc, entry["name"], level=3)
                doc.add_paragraph(entry.get("summary", ""))

    # ── Recommendations Summary ──────────────────────────────────────────
    _add_heading(doc, "Recommendations Summary", level=1)
    doc.add_paragraph(data.get("recommendations_summary", ""))

    # ── Conclusion ───────────────────────────────────────────────────────
    _add_heading(doc, "Conclusion", level=1)
    doc.add_paragraph(data.get("conclusion", ""))

    # ── Certification Methodology Note ───────────────────────────────────
    if data.get("certification_methodology"):
        _add_heading(doc, "Certification Methodology Note", level=2)
        doc.add_paragraph(data["certification_methodology"])

    # ── Verification Sources ─────────────────────────────────────────────
    if data.get("verification_sources"):
        _add_heading(doc, "Verification Sources", level=2)
        for src in data["verification_sources"]:
            doc.add_paragraph(f"{src['vendor']}: {src['url']}", style="List Bullet")

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── CLI entry point ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: generate_compliance_matrix.py <json_file> <output_path>")
        sys.exit(1)

    json_path = Path(sys.argv[1])
    output = sys.argv[2]

    with open(json_path) as f:
        data = json.load(f)

    generate_compliance_matrix(data, output)
