"""Generate a Supplier Selection Questionnaire .docx document.

Usage:
    uv run --package elt-doc-vendor-assess python elt_doc_vendor_assess/src/elt_doc_vendor_assess/generate_supplier_selection.py <json_file> <output_path>

Input JSON structure:
{
    "vendor_name": "Hyve Managed Hosting",
    "product_name": "Cerberus SFTP Server",
    "document_history": [
        {"version": "1.0", "date": "2026-02-10", "author": "Rakesh Patel", "summary": "Initial version"}
    ],
    "reference_documents": [
        {"document": "SAD Document", "link": "..."}
    ],
    "sections": {
        "functionality": {
            "core_functionality": "...",
            "use_case_diagram": "...",
            "architecture_diagram": "..."
        },
        "extension": {
            "customization": "..."
        },
        "data": {
            "data_held": "...",
            "data_classification": "...",
            "data_protection": "...",
            "testing_results": "..."
        },
        "ai": {
            "model_training": "...",
            "data_leaving_tenant": "...",
            "external_ai_vendors": "...",
            "bias_mitigation": "...",
            "gdpr_ai": "...",
            "eu_ai_act": "...",
            "ai_transparency": "...",
            "ai_human_oversight": "...",
            "ai_data_quality": "...",
            "ai_risk_assessment": "...",
            "ai_documentation": "..."
        },
        "security": {
            "saas_hosting": "...",
            "shared_dedicated": "...",
            "pen_testing": "...",
            "vulnerability_scans": "...",
            "patching": "...",
            "authentication": "...",
            "admin_access": "...",
            "vendor_data_access": "...",
            "encryption": "...",
            "intrusion_detection": "...",
            "breach_history": "...",
            "security_frameworks": "...",
            "high_availability": "...",
            "backup": "..."
        },
        "performance": {
            "scalability": "...",
            "benchmarks": "...",
            "slas": "...",
            "sla_enforcement": "..."
        },
        "compliance": {
            "regulatory": "...",
            "certifications": "..."
        },
        "reliability": {
            "incident_history": "...",
            "observability": "...",
            "dr_plan": "...",
            "rpo": "...",
            "rto": "...",
            "failure_testing": "...",
            "integration_reliability": "..."
        },
        "hosting": {
            "location": "...",
            "infrastructure": "...",
            "environments": "..."
        },
        "technology": {
            "tech_stack": "...",
            "patching_strategy": "...",
            "release_strategy": "...",
            "breaking_changes": "...",
            "vulnerability_report": "..."
        },
        "integration_api": {
            "integration_options": "...",
            "documentation": "...",
            "data_migration": "..."
        },
        "roadmap": {
            "planned_releases": "..."
        },
        "appendix": "..."
    }
}
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ── Section definitions ──────────────────────────────────────────────────────

SECTIONS = [
    ("1. Functionality", "functionality", [
        ("1.1", "Core Functionality Overview", "core_functionality",
         "Provide an overview of the core features and functionality of the product."),
        ("1.2", "Use Case Diagram", "use_case_diagram",
         "Provide a use case diagram showing key user interactions."),
        ("1.3", "Architecture Context", "architecture_diagram",
         "Provide an L1 architecture context diagram showing how the product fits within the FA technology landscape."),
    ]),
    ("2. Extension", "extension", [
        ("2.1", "Customization Capabilities", "customization",
         "How can the product be customized and integrated with other systems?"),
    ]),
    ("3. Data", "data", [
        ("3.1", "Data Held", "data_held",
         "What data does the product hold or process?"),
        ("3.2", "Data Classification", "data_classification",
         "How is data classified within the system?"),
        ("3.3", "Data Protection", "data_protection",
         "What data protection measures are in place?"),
        ("3.4", "Testing Results", "testing_results",
         "Provide any relevant testing results or certifications."),
    ]),
    ("4. AI", "ai", [
        ("4.1", "Model Training", "model_training",
         "Does the product use AI/ML models? How are they trained?"),
        ("4.2", "Data Leaving Tenant", "data_leaving_tenant",
         "Does any data leave the tenant for AI processing?"),
        ("4.3", "External AI Vendors", "external_ai_vendors",
         "Are any external AI vendors or services used?"),
        ("4.4", "Bias Mitigation", "bias_mitigation",
         "What measures are in place to mitigate AI bias?"),
        ("4.5", "GDPR & AI", "gdpr_ai",
         "How does the product comply with GDPR in relation to AI features?"),
        ("4.6", "EU AI Act", "eu_ai_act",
         "How does the product comply with the EU AI Act?"),
        ("4.7", "AI Transparency", "ai_transparency",
         "How is AI transparency ensured for end users?"),
        ("4.8", "Human Oversight", "ai_human_oversight",
         "What human oversight mechanisms exist for AI decisions?"),
        ("4.9", "Data Quality for AI", "ai_data_quality",
         "How is data quality maintained for AI training and inference?"),
        ("4.10", "Risk Assessment", "ai_risk_assessment",
         "What risk assessment has been performed on AI capabilities?"),
        ("4.11", "AI Documentation", "ai_documentation",
         "What documentation is available for AI features and their limitations?"),
    ]),
    ("5. Security", "security", [
        ("5.1", "SaaS Hosting", "saas_hosting",
         "Is the product SaaS-hosted? Describe the hosting model."),
        ("5.2", "Shared or Dedicated", "shared_dedicated",
         "Is the environment shared or dedicated?"),
        ("5.3", "Penetration Testing", "pen_testing",
         "How frequently are penetration tests conducted? Share recent results."),
        ("5.4", "Vulnerability Scanning", "vulnerability_scans",
         "Describe the vulnerability scanning process and frequency."),
        ("5.5", "Patching", "patching",
         "What is the patching strategy and cadence?"),
        ("5.6", "Authentication", "authentication",
         "What authentication mechanisms are supported?"),
        ("5.7", "Admin Access", "admin_access",
         "How is administrative access managed and audited?"),
        ("5.8", "Vendor Data Access", "vendor_data_access",
         "What access does the vendor have to customer data?"),
        ("5.9", "Encryption", "encryption",
         "Describe encryption at rest and in transit."),
        ("5.10", "Intrusion Detection", "intrusion_detection",
         "What intrusion detection/prevention systems are in place?"),
        ("5.11", "Breach History", "breach_history",
         "Has there been any security breach in the last 5 years?"),
        ("5.12", "Security Frameworks", "security_frameworks",
         "Which security frameworks or standards does the product comply with?"),
        ("5.13", "High Availability", "high_availability",
         "Describe the high availability architecture."),
        ("5.14", "Backup & Recovery", "backup",
         "Describe backup procedures, retention, and recovery processes."),
    ]),
    ("6. Performance", "performance", [
        ("6.1", "Scalability", "scalability",
         "How does the product scale to meet growing demand?"),
        ("6.2", "Performance Benchmarks", "benchmarks",
         "Provide performance benchmarks or test results."),
        ("6.3", "SLAs", "slas",
         "What SLAs are offered for availability and performance?"),
        ("6.4", "SLA Enforcement", "sla_enforcement",
         "How are SLAs monitored and enforced?"),
    ]),
    ("7. Compliance and Regulations", "compliance", [
        ("7.1", "Regulatory Compliance", "regulatory",
         "What regulatory requirements does the product comply with?"),
        ("7.2", "Certifications", "certifications",
         "List all relevant certifications held."),
    ]),
    ("8. Reliability", "reliability", [
        ("8.1", "Incident History", "incident_history",
         "Provide details of any significant incidents in the past 2 years."),
        ("8.2", "Observability", "observability",
         "What monitoring and observability tools are used?"),
        ("8.3", "Disaster Recovery Plan", "dr_plan",
         "Describe the disaster recovery plan."),
        ("8.4", "RPO", "rpo",
         "What is the Recovery Point Objective?"),
        ("8.5", "RTO", "rto",
         "What is the Recovery Time Objective?"),
        ("8.6", "Failure Testing", "failure_testing",
         "How is failure testing conducted (chaos engineering, etc.)?"),
        ("8.7", "Integration Reliability", "integration_reliability",
         "How is reliability ensured for integrations with external systems?"),
    ]),
    ("9. Hosting", "hosting", [
        ("9.1", "Hosting Location", "location",
         "Where is the product hosted? Specify data centre locations."),
        ("9.2", "Infrastructure Details", "infrastructure",
         "Describe the underlying infrastructure."),
        ("9.3", "Environments", "environments",
         "What environments are available (dev, staging, production)?"),
    ]),
    ("10. Technology", "technology", [
        ("10.1", "Technology Stack", "tech_stack",
         "Describe the technology stack used."),
        ("10.2", "Patching Strategy", "patching_strategy",
         "What is the patching and update strategy?"),
        ("10.3", "Release Strategy", "release_strategy",
         "Describe the release cadence and deployment strategy."),
        ("10.4", "Breaking Changes", "breaking_changes",
         "How are breaking changes communicated and managed?"),
        ("10.5", "Vulnerability Reporting", "vulnerability_report",
         "How are vulnerabilities reported and tracked?"),
    ]),
    ("11. Integration and API", "integration_api", [
        ("11.1", "Integration Options", "integration_options",
         "What integration options are available (APIs, webhooks, file-based)?"),
        ("11.2", "Documentation", "documentation",
         "Where can API/integration documentation be found?"),
        ("11.3", "Data Migration", "data_migration",
         "What data migration tools or support are available?"),
    ]),
    ("12. Roadmap", "roadmap", [
        ("12.1", "Planned Releases", "planned_releases",
         "Describe planned releases and future development direction."),
    ]),
    ("13. Appendix", "appendix", []),
]


# ── Styling helpers ──────────────────────────────────────────────────────────

def _set_cell_text(cell, text, bold=False, font_size=9):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Calibri"


def _shade_cell(cell, color_hex):
    """Apply background shading to a cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _style_header_row(row, bg_color="1F4E79"):
    """Style a table header row."""
    for cell in row.cells:
        _shade_cell(cell, bg_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True


def _add_heading(doc, text, level):
    """Add a heading with consistent font."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


# ── Main document builder ────────────────────────────────────────────────────

def generate_supplier_selection(data: dict, output_path: str):
    """Generate the Supplier Selection Questionnaire .docx document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title ────────────────────────────────────────────────────────────
    title = f"Supplier Selection Questionnaire — {data['vendor_name']}"
    if data.get("product_name"):
        title += f" ({data['product_name']})"
    _add_heading(doc, title, level=0)

    # ── Document History Table ───────────────────────────────────────────
    _add_heading(doc, "Document History", level=2)
    history = data.get("document_history", [])
    if history:
        cols = ["Version No", "Revision Date", "Author", "Summary of Changes"]
        table = doc.add_table(rows=1 + len(history), cols=len(cols))
        table.style = "Table Grid"
        for i, col in enumerate(cols):
            _set_cell_text(table.rows[0].cells[i], col, bold=True)
        _style_header_row(table.rows[0])
        for r, h in enumerate(history, start=1):
            _set_cell_text(table.rows[r].cells[0], h.get("version", ""))
            _set_cell_text(table.rows[r].cells[1], h.get("date", ""))
            _set_cell_text(table.rows[r].cells[2], h.get("author", ""))
            _set_cell_text(table.rows[r].cells[3], h.get("summary", ""))
    doc.add_paragraph()

    # ── Reference Documents Table ────────────────────────────────────────
    refs = data.get("reference_documents", [])
    if refs:
        _add_heading(doc, "Reference Documents", level=2)
        table = doc.add_table(rows=1 + len(refs), cols=2)
        table.style = "Table Grid"
        _set_cell_text(table.rows[0].cells[0], "Document", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Link", bold=True)
        _style_header_row(table.rows[0])
        for r, ref in enumerate(refs, start=1):
            _set_cell_text(table.rows[r].cells[0], ref.get("document", ""))
            _set_cell_text(table.rows[r].cells[1], ref.get("link", ""))
        doc.add_paragraph()

    # ── Sections ─────────────────────────────────────────────────────────
    sections_data = data.get("sections", {})

    for section_heading, section_key, sub_questions in SECTIONS:
        _add_heading(doc, section_heading, level=1)
        section_content = sections_data.get(section_key, {})

        if section_key == "appendix":
            # Appendix is freeform text
            if isinstance(section_content, str) and section_content:
                doc.add_paragraph(section_content)
            elif isinstance(section_content, dict):
                for key, val in section_content.items():
                    doc.add_paragraph(str(val))
            continue

        if isinstance(section_content, str):
            # Single answer for the whole section
            doc.add_paragraph(section_content)
            continue

        for num, sub_heading, field_key, question_text in sub_questions:
            _add_heading(doc, f"{num} {sub_heading}", level=2)
            answer = section_content.get(field_key, "")
            if answer:
                # Split multi-paragraph answers
                for para in str(answer).split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())
            else:
                doc.add_paragraph(f"[{question_text}]")

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Saved: {output_path}")


# ── CLI entry point ──────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: generate_supplier_selection.py <json_file> <output_path>")
        sys.exit(1)

    json_path = Path(sys.argv[1])
    output = sys.argv[2]

    with open(json_path) as f:
        data = json.load(f)

    generate_supplier_selection(data, output)
