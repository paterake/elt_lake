"""Analyze detailed styling and table structure of two compliance matrix documents."""

from pathlib import Path
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    None: "INHERITED/DEFAULT",
}


def get_alignment_name(alignment):
    return ALIGN_MAP.get(alignment, f"UNKNOWN({alignment})")


def analyze_run_formatting(run):
    """Extract formatting details from a run."""
    info = {}
    if run.bold:
        info["bold"] = True
    if run.italic:
        info["italic"] = True
    if run.underline:
        info["underline"] = True
    if run.font.size:
        info["font_size"] = f"{run.font.size.pt}pt"
    if run.font.name:
        info["font_name"] = run.font.name
    if run.font.color and run.font.color.rgb:
        info["font_color"] = str(run.font.color.rgb)
    if run.font.highlight_color:
        info["highlight"] = str(run.font.highlight_color)
    return info


def analyze_document(filepath):
    """Full analysis of a single docx document."""
    doc = Document(filepath)
    filename = Path(filepath).name

    print("=" * 100)
    print(f"  DOCUMENT: {filename}")
    print("=" * 100)

    # -- Section properties --
    for i, section in enumerate(doc.sections):
        print(f"\n--- Section {i} Properties ---")
        print(f"  Page width:  {section.page_width.inches:.2f} in")
        print(f"  Page height: {section.page_height.inches:.2f} in")
        print(f"  Left margin: {section.left_margin.inches:.2f} in")
        print(f"  Right margin:{section.right_margin.inches:.2f} in")
        print(f"  Top margin:  {section.top_margin.inches:.2f} in")
        print(f"  Bot margin:  {section.bottom_margin.inches:.2f} in")
        if section.orientation is not None:
            print(f"  Orientation: {section.orientation}")

    # -- Paragraph styles --
    print(f"\n{'~' * 80}")
    print("  PARAGRAPH STYLES USED")
    print(f"{'~' * 80}")

    style_counter = Counter()
    heading_fonts = {}
    para_details = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"
        style_counter[style_name] += 1

        # Collect heading font sizes
        if "Heading" in style_name or "Title" in style_name:
            for run in para.runs:
                if run.font.size:
                    heading_fonts[style_name] = run.font.size.pt

        # Collect first ~60 paragraphs with details
        if len(para_details) < 60:
            detail = {
                "style": style_name,
                "alignment": get_alignment_name(para.alignment),
                "text_preview": para.text[:120] if para.text else "(empty)",
            }
            if para.runs:
                detail["run_formats"] = [
                    analyze_run_formatting(r) for r in para.runs[:3]
                ]
            para_details.append(detail)

    print("\n  Style Name                          Count")
    print("  " + "-" * 50)
    for style_name, count in style_counter.most_common():
        print(f"  {style_name:<38} {count}")

    # -- Heading font sizes --
    print(f"\n{'~' * 80}")
    print("  HEADING / TITLE FONT SIZES (from runs)")
    print(f"{'~' * 80}")

    for style_name in sorted(style_counter.keys()):
        if "Heading" in style_name or "Title" in style_name:
            style_obj = None
            for s in doc.styles:
                if s.name == style_name:
                    style_obj = s
                    break
            run_size = heading_fonts.get(style_name, "not found in runs")
            style_font_size = "N/A"
            style_font_name = "N/A"
            if style_obj and style_obj.font:
                if style_obj.font.size:
                    style_font_size = f"{style_obj.font.size.pt}pt"
                if style_obj.font.name:
                    style_font_name = style_obj.font.name
            print(f"  {style_name:<30} run_size={str(run_size):<12} style_size={style_font_size:<10} style_font={style_font_name}")

    # -- Paragraph details --
    print(f"\n{'~' * 80}")
    print("  PARAGRAPH DETAILS (first ~60 paragraphs)")
    print(f"{'~' * 80}")

    for i, d in enumerate(para_details):
        text_preview = d["text_preview"]
        if not text_preview.strip() and text_preview == "(empty)":
            continue
        print(f"\n  [{i:03d}] Style: {d['style']:<30} Align: {d['alignment']}")
        print(f"        Text: {text_preview}")
        if "run_formats" in d:
            for ri, rf in enumerate(d["run_formats"]):
                if rf:
                    print(f"        Run[{ri}] formatting: {rf}")

    # -- Bullet / list detection --
    print(f"\n{'~' * 80}")
    print("  BULLET / LIST STYLE DETECTION")
    print(f"{'~' * 80}")

    bullet_styles = defaultdict(list)
    for para in doc.paragraphs:
        sname = para.style.name if para.style else "None"
        if "list" in sname.lower() or "bullet" in sname.lower():
            bullet_styles[sname].append(para.text[:80])
        pPr = para._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        if pPr is not None:
            ilvl_el = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl")
            numId_el = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
            ilvl = ilvl_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if ilvl_el is not None else "?"
            numId = numId_el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") if numId_el is not None else "?"
            key = f"{sname} [numId={numId}, ilvl={ilvl}]"
            if len(bullet_styles[key]) < 3:
                bullet_styles[key].append(para.text[:80])

    if bullet_styles:
        for key, examples in bullet_styles.items():
            print(f"\n  {key}:")
            for ex in examples[:3]:
                print(f"    - {ex}")
    else:
        print("  (no bullet/list paragraphs detected)")

    # -- Tables --
    print(f"\n{'~' * 80}")
    print(f"  TABLES ({len(doc.tables)} found)")
    print(f"{'~' * 80}")

    for ti, table in enumerate(doc.tables):
        rows = table.rows
        cols = table.columns
        print(f"\n  +-- TABLE {ti} -----------------------------------------------")
        print(f"  | Rows: {len(rows)}, Columns: {len(cols)}")

        if table.style:
            print(f"  | Table style: {table.style.name}")

        # Column widths
        try:
            widths = []
            for cell in rows[0].cells:
                w = cell.width
                if w:
                    widths.append(f"{w.inches:.2f}in")
                else:
                    widths.append("auto")
            print(f"  | Column widths (row 0): {widths}")
        except Exception:
            pass

        # Header row
        print(f"  |")
        print(f"  | HEADER ROW (row 0):")
        for ci, cell in enumerate(rows[0].cells):
            cell_text = cell.text.strip().replace("\n", " | ")
            alignment = "N/A"
            bold = False
            italic = False
            font_size = "N/A"
            font_name = "N/A"
            bg_color = "N/A"

            if cell.paragraphs:
                p = cell.paragraphs[0]
                alignment = get_alignment_name(p.alignment)
                for run in p.runs:
                    if run.bold:
                        bold = True
                    if run.italic:
                        italic = True
                    if run.font.size:
                        font_size = f"{run.font.size.pt}pt"
                    if run.font.name:
                        font_name = run.font.name

            tc = cell._tc
            tcPr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
            if tcPr is not None:
                shd = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                if shd is not None:
                    bg_color = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "N/A")

            print(f"  |   Col[{ci}]: \"{cell_text}\"")
            print(f"  |          align={alignment}, bold={bold}, italic={italic}, font={font_name}/{font_size}, bg={bg_color}")

        # Data rows (first 3)
        max_data_rows = min(3, len(rows) - 1)
        for ri in range(1, 1 + max_data_rows):
            row = rows[ri]
            print(f"  |")
            print(f"  | DATA ROW {ri}:")
            for ci, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace("\n", " | ")[:150]
                alignment = "N/A"
                bold = False
                italic = False
                font_size = "N/A"
                font_name = "N/A"
                bg_color = "N/A"

                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    alignment = get_alignment_name(p.alignment)
                    for run in p.runs:
                        if run.bold:
                            bold = True
                        if run.italic:
                            italic = True
                        if run.font.size:
                            font_size = f"{run.font.size.pt}pt"
                        if run.font.name:
                            font_name = run.font.name

                tc = cell._tc
                tcPr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
                if tcPr is not None:
                    shd = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                    if shd is not None:
                        fill = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "N/A")
                        bg_color = fill

                gridSpan = None
                if tcPr is not None:
                    gs = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan")
                    if gs is not None:
                        gridSpan = gs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")

                merge_note = f" [MERGED span={gridSpan}]" if gridSpan and int(gridSpan) > 1 else ""
                fmt_parts = []
                if bold:
                    fmt_parts.append("BOLD")
                if italic:
                    fmt_parts.append("ITALIC")
                fmt_str = ",".join(fmt_parts) if fmt_parts else "normal"

                print(f"  |   Col[{ci}]: \"{cell_text}\"{merge_note}")
                print(f"  |          align={alignment}, fmt={fmt_str}, font={font_name}/{font_size}, bg={bg_color}")

        # Merge/span analysis
        print(f"  |")
        print(f"  | MERGE / SPAN ANALYSIS:")
        merge_info = []
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr")
                if tcPr is not None:
                    gs = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan")
                    vmerge = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge")
                    if gs is not None:
                        val = gs.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1")
                        if int(val) > 1:
                            merge_info.append(f"    Row {ri}, Col {ci}: horizontal span={val}")
                    if vmerge is not None:
                        val = vmerge.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "continue")
                        merge_info.append(f"    Row {ri}, Col {ci}: vertical merge={val}")
        if merge_info:
            for m in merge_info[:20]:
                print(f"  | {m}")
            if len(merge_info) > 20:
                print(f"  |   ... and {len(merge_info) - 20} more merges")
        else:
            print(f"  |   (no merged cells)")

        print(f"  +-------------------------------------------------------------")

    # -- Formatting patterns summary --
    print(f"\n{'~' * 80}")
    print("  FORMATTING PATTERNS SUMMARY")
    print(f"{'~' * 80}")

    bold_paras = []
    for para in doc.paragraphs:
        if para.runs and all(r.bold for r in para.runs if r.text.strip()):
            if para.text.strip():
                bold_paras.append((para.style.name, para.text[:80]))
    print(f"\n  All-bold paragraphs ({len(bold_paras)}):")
    for sn, txt in bold_paras[:10]:
        print(f"    [{sn}] {txt}")

    all_sizes = Counter()
    all_fonts = Counter()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                all_sizes[run.font.size.pt] += 1
            if run.font.name:
                all_fonts[run.font.name] += 1

    print(f"\n  Font sizes used across document:")
    for size, count in all_sizes.most_common():
        print(f"    {size}pt -- {count} occurrences")

    print(f"\n  Font names used across document:")
    for name, count in all_fonts.most_common():
        print(f"    {name} -- {count} occurrences")

    print("\n")


# -- Main --
if __name__ == "__main__":
    docs = [
        Path.home() / "Downloads" / "CRM_Migration_Tool_Compliance_Assessment.docx",
        Path.home() / "Downloads" / "SaaS_SFTP_Compliance_Matrix.docx",
    ]
    for docpath in docs:
        analyze_document(str(docpath))
