#!/usr/bin/env python3
"""
Generate Workday Integration XML from SAD Document
"""
import sys
import re
import os
import argparse
import pandas as pd
from pathlib import Path
from docx import Document
from elt_doc_sad_leanix.diagram_generator import WorkdayIntegrationDiagramGenerator

def load_inventory(inventory_path):
    """Load LeanIX inventory from Excel"""
    try:
        df = pd.read_excel(inventory_path)
        # Ensure we have required columns
        required = ['id', 'type', 'name']
        if not all(col in df.columns for col in required):
            print(f"Warning: Inventory missing required columns: {required}")
            return None
        return df
    except Exception as e:
        print(f"Error loading inventory: {e}")
        return None

def lookup_factsheet(df, name, type_filter=None):
    """Find FactSheet ID and Type by name"""
    if df is None or not name:
        return None, None
    
    # Normalize name for search
    name_lower = name.lower().strip()
    
    # 1. Exact match
    match = df[df['name'].str.lower() == name_lower]
    
    # 2. Contains match (if no exact match)
    if match.empty:
        match = df[df['name'].str.lower().str.contains(name_lower, regex=False)]
        
    # Filter by type if provided
    if not match.empty and type_filter:
        match = match[match['type'] == type_filter]
        
    if not match.empty:
        # Return first match
        row = match.iloc[0]
        return row['id'], row['type']
        
    return None, None

def parse_sad(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    spec = {
        'title': '',
        'integration_id': '',
        'source_system': 'Workday Human Capital Management',
        'intermediary': '',
        'target_system': '',
        'direction': 'outbound', # Default
        'flow_labels': [],
        'security_details': [],
        'system_of_record': [],
        'key_attributes': [],
        'notes': []
    }
    
    # 1. Integration ID
    int_id_match = re.search(r'(INT\d{3})', full_text)
    if int_id_match:
        spec['integration_id'] = int_id_match.group(1)
        
    # 2. Target System (Simple heuristic: usually the name after "INTxxx" or in title)
    # Looking for "Target | <Name>" in tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                if 'Target' in cells[0] and 'Technology' in cells[0]:
                     pass # Header
                elif cells[0] == 'Target':
                    spec['target_system'] = cells[1].split('(')[0].strip()
                    if "Health & Safety Platform" in spec['target_system']:
                         spec['target_system'] = "Cardinus" # Cleanup
                elif cells[0] == 'Source':
                    if 'Workday' in cells[1]:
                        if 'Financial' in cells[1]:
                            spec['source_system'] = 'Workday Financial Management'
                        else:
                            spec['source_system'] = 'Workday Human Capital Management'
                elif cells[0] == 'SFTP Service' or cells[0] == 'File Transfer':
                     if 'SFTP' in cells[1]:
                         spec['intermediary'] = cells[1]

    # Refine Intermediary
    if not spec['intermediary'] and 'SFTP' in full_text:
        # Check if Cardinus SFTP is mentioned
        if 'Cardinus Managed SFTP' in full_text:
            spec['intermediary'] = 'Cardinus SFTP'
        elif 'FA Managed SFTP' in full_text:
            spec['intermediary'] = 'FA SFTP'

    # Refine Target if not found
    if not spec['target_system']:
        if 'Cardinus' in full_text:
            spec['target_system'] = 'Cardinus'

    if not spec['target_system']:
        stem = Path(docx_path).stem
        parts = stem.split('_')
        if len(parts) >= 3:
            core = parts[2:]
            if len(core) >= 2 and core[-2].startswith('V'):
                core = core[:-2]
            candidate = " ".join(core).strip()
            if candidate:
                spec['target_system'] = candidate
        if not spec['target_system']:
            spec['target_system'] = 'Target System'

    if spec['integration_id'] == 'INT018':
        spec['direction'] = 'bidirectional'

    spec['title'] = f"Workday {spec['target_system']} Integration"
    
    # Direction
    if 'Outbound' in full_text or 'outbound' in full_text:
        spec['direction'] = 'outbound'
    elif 'Inbound' in full_text or 'inbound' in full_text:
        spec['direction'] = 'inbound'
        
    # Flow Labels (Heuristic extraction from Functionality section)
    # We'll construct generic ones based on found components
    if spec['direction'] == 'outbound':
        target_label = spec['target_system'] or 'Target System'
        spec['flow_labels'] = [
            {
                'text': f"Workday EIB extracts data via CR {spec['integration_id']}",
                'x': 100, 'y': 180, 'width': 300, 'height': 60
            },
            {
                'text': f"PGP Encrypted & Uploaded to {spec['intermediary'] or 'SFTP'}",
                'x': 450, 'y': 180, 'width': 300, 'height': 60
            },
             {
                'text': f"{target_label} imports data",
                'x': 800, 'y': 180, 'width': 300, 'height': 60
            }
        ]
    elif spec['direction'] == 'inbound':
        spec['flow_labels'] = [
            {
                'text': f"{spec['target_system']} sends data (encrypted)",
                'x': 100, 'y': 180, 'width': 300, 'height': 60
            },
            {
                'text': f"File placed on {spec['intermediary'] or 'SFTP'}",
                'x': 450, 'y': 180, 'width': 300, 'height': 60
            },
             {
                'text': f"Workday retrieves and processes file",
                'x': 800, 'y': 180, 'width': 300, 'height': 60
            }
        ]
    
    # Security Details and Component Checks
    for table in doc.tables:
        if len(table.rows) < 1:
            continue
            
        headers = [c.text.strip().lower() for c in table.rows[0].cells]

        # Check for Component/Technology tables
        if "component" in headers and "technology" in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    comp = cells[0].lower()
                    tech = cells[1]
                    if "data source" in comp:
                        spec["process_extraction"] = tech
                    elif "integration development" in comp:
                        spec["process_processing"] = tech

        # Extract from "Control | Implementation" table
        if 'control' in headers and 'implementation' in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    spec['security_details'].append(f"{cells[0]}: {cells[1]}")
    
    # Key Attributes
    # Extract from field mapping table
    for table in doc.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        if 'cardinus field' in headers or 'target field' in headers:
            for row in table.rows[1:]:
                cells = [c.text.strip() for c in row.cells]
                if cells[0]:
                    spec['key_attributes'].append(cells[0])
    
    # Notes / Assumptions
    # Look for headers like "Notes", "Assumptions", "Constraints"
    capturing_notes = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.lower() in ['notes', 'assumptions', 'constraints', 'additional notes']:
            capturing_notes = True
            continue
        
        if capturing_notes:
            # Stop if we hit a new major header (assuming headers are short and bold/large, 
            # but here we just check for empty or specific keywords if needed, 
            # or just take the next few paragraphs)
            if not text:
                continue
            if len(text) < 30 and text.isupper(): # heuristic for next header
                capturing_notes = False
                continue
            spec['notes'].append(text)

    # System of Record
    spec['system_of_record'] = [
        f"{spec['source_system']}: Employee Master Data",
        f"{spec['target_system']}: Health & Safety Training"
    ]
    
    # Process Steps Defaults
    if 'process_extraction' not in spec:
        spec['process_extraction'] = "Workday Custom Report"
        
    if 'process_processing' not in spec:
        spec['process_processing'] = "Workday Studio / EIB"
        
    # Process Transmission
    if spec.get('intermediary'):
        spec['process_transmission'] = f"Via {spec['intermediary']} (SFTP)"
    else:
        spec['process_transmission'] = "Direct API / HTTP"
        
    # Process Security
    if spec['security_details']:
        # Summarize security
        sec_summary = []
        for item in spec['security_details']:
            item_lower = item.lower()
            if "pgp" in item_lower or "encryption" in item_lower or "tls" in item_lower or "ssl" in item_lower:
                # Extract value after colon if present
                val = item.split(':', 1)[1].strip() if ':' in item else item
                sec_summary.append(val)
        spec['process_security'] = "; ".join(sec_summary) if sec_summary else "Standard Encryption"
    else:
        spec['process_security'] = "Standard Encryption"

    return spec

def main():
    parser = argparse.ArgumentParser(description='Generate LeanIX XML from SAD')
    parser.add_argument('docx_path', help='Path to SAD docx file')
    args = parser.parse_args()
    
    docx_path = Path(args.docx_path)
    if not docx_path.exists():
        print(f"Error: File {docx_path} not found")
        sys.exit(1)
    
    # Load Inventory
    inventory_path = Path("config/LeanIX_Inventory.xlsx")
    inventory_df = None
    if inventory_path.exists():
        print(f"Loading inventory from {inventory_path}...")
        inventory_df = load_inventory(inventory_path)
    else:
        print(f"Warning: Inventory file not found at {inventory_path}")
        
    print(f"Reading {docx_path}...")
    spec = parse_sad(docx_path)
    
    # Lookup IDs
    if inventory_df is not None:
        # Source
        sid, stype = lookup_factsheet(inventory_df, spec['source_system'], "Application")
        if sid:
            spec['source_id'] = sid
            spec['source_type'] = stype
            print(f"Found Source: {spec['source_system']} -> {sid}")
            
        # Target
        tid, ttype = lookup_factsheet(inventory_df, spec['target_system'])
        if tid:
            spec['target_id'] = tid
            spec['target_type'] = ttype
            print(f"Found Target: {spec['target_system']} -> {tid}")
            
        # Intermediary
        if spec['intermediary']:
            iid, itype = lookup_factsheet(inventory_df, spec['intermediary'])
            if not iid and 'SFTP' in spec['intermediary']:
                # Fallback for generic SFTP
                iid, itype = lookup_factsheet(inventory_df, "SFTP", "ITComponent")
            
            if iid:
                spec['intermediary_id'] = iid
                spec['intermediary_type'] = itype
                print(f"Found Intermediary: {spec['intermediary']} -> {iid}")

        # Interface (Edge)
        if spec['integration_id']:
            # Search for Interface by INT ID (e.g. "INT011")
            int_id = spec['integration_id']
            if_id, if_type = lookup_factsheet(inventory_df, int_id, "Interface")
            
            if if_id:
                spec['interface_id'] = if_id
                spec['interface_label'] = f"{spec['source_system']} - {spec['target_system']}" # Default label
                # Try to get actual name from inventory if possible, but lookup returns ID/Type. 
                # We could modify lookup to return name too, but this is fine for now.
                print(f"Found Interface: {int_id} -> {if_id}")
    
    print("Extracted Spec:")
    for k, v in spec.items():
        if k not in ['flow_labels', 'security_details', 'system_of_record', 'key_attributes', 'notes']:
            print(f"{k}: {v}")
        
    generator = WorkdayIntegrationDiagramGenerator()
    xml = generator.generate_xml(spec)
    
    output_path = docx_path.with_suffix('.xml')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(xml)
        
    print(f"✅ Generated: {output_path}")

if __name__ == '__main__':
    main()
