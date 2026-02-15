#!/usr/bin/env python3
"""
Compile a full LLM prompt by combining:
1. The Master Prompt (Instructions)
2. The LeanIX Inventory (Context)
3. The SAD Document (Input)
"""
import sys
import argparse
import openpyxl
from pathlib import Path
from docx import Document

def extract_inventory(path):
    """Convert inventory Excel to Markdown table using openpyxl"""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        rows = list(ws.iter_rows(values_only=True))
        wb.close()

        if len(rows) < 2:
            return "Error: Inventory has no data rows"

        # Row 0 = machine keys (id, type, name, ...), row 1 = human labels, data from row 2
        headers = [str(h) if h else "" for h in rows[0]]

        # Find column indices for name, type, id
        col_map = {}
        for i, h in enumerate(headers):
            h_lower = h.lower().strip()
            if h_lower == "name":
                col_map["name"] = i
            elif h_lower == "type":
                col_map["type"] = i
            elif h_lower == "id":
                col_map["id"] = i

        required = ["name", "type", "id"]
        missing = [c for c in required if c not in col_map]
        if missing:
            return f"Error: Inventory missing required columns: {missing}"

        # Build markdown table
        lines = [
            "| name | type | id |",
            "|------|------|----|",
        ]
        for row in rows[2:]:  # skip both header rows
            name = row[col_map["name"]] or ""
            typ = row[col_map["type"]] or ""
            rid = row[col_map["id"]] or ""
            if name:  # skip empty rows
                lines.append(f"| {name} | {typ} | {rid} |")

        return "\n".join(lines)
    except Exception as e:
        return f"Error reading inventory: {e}"

def extract_sad_text(path):
    """Extract text from SAD docx"""
    try:
        doc = Document(path)
        full_text = []

        # Add paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text)

        # Add tables (linearized)
        for table in doc.tables:
            full_text.append("\n[TABLE START]")
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(" | ".join(cells))
            full_text.append("\n".join(rows))
            full_text.append("[TABLE END]\n")

        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading SAD: {e}"

def main():
    parser = argparse.ArgumentParser(description='Compile LLM Prompt')
    # Determine default paths relative to this script
    script_dir = Path(__file__).parent
    # Prompts are in ../prompts/ relative to this script (src/elt_doc_sad_leanix/cmd -> src/elt_doc_sad_leanix/prompts)
    default_prompt = script_dir.parent / 'prompts' / 'sad_to_leanix.md'
    # script_dir is src/elt_doc_sad_leanix/cmd
    project_root = script_dir.parent.parent.parent.parent
    if not (project_root / 'pyproject.toml').exists():
         project_root = script_dir.parent.parent.parent
    default_inventory = project_root / 'config' / 'LeanIX_Inventory.xlsx'
    if not default_inventory.exists():
         default_inventory = project_root / 'elt_doc_sad_leanix' / 'config' / 'LeanIX_Inventory.xlsx'

    parser.add_argument('sad_path', help='Path to SAD docx file')
    parser.add_argument('--prompt', default=str(default_prompt), help='Path to master prompt file')
    parser.add_argument('--inventory', default=str(default_inventory), help='Path to inventory file')
    args = parser.parse_args()

    # Read files
    prompt_path = Path(args.prompt)
    inventory_path = Path(args.inventory)
    sad_path = Path(args.sad_path)

    if not prompt_path.exists():
        print(f"Error: Prompt file {prompt_path} not found")
        sys.exit(1)

    print(f"Compiling prompt for {sad_path}...")

    with open(prompt_path, 'r', encoding='utf-8') as f:
        master_prompt = f.read()

    inventory_text = extract_inventory(inventory_path)
    sad_text = extract_sad_text(sad_path)

    # List templates
    templates_dir = project_root / 'elt_doc_sad_leanix' / 'templates'
    if not templates_dir.exists():
         templates_dir = project_root / 'templates'

    template_list = "No templates found."
    if templates_dir.exists():
        templates = [f.name for f in templates_dir.glob('*.xml')]
        template_list = "\n".join([f"* {t}" for t in sorted(templates)])

    # Assemble
    full_output = f"""{master_prompt}

---

## CONTEXT DATA

### Available Templates
The following XML templates are available in the repository. You MUST select the one that best matches the integration pattern.

{template_list}

### LeanIX Inventory
The following is the authoritative list of systems and IDs. You MUST reference these IDs.

{inventory_text}

---

## INPUT DATA

### SAD Document Content
The following is the extracted text from the System Architecture Document.

{sad_text}
"""

    # Output to stdout (or file)
    print(full_output)

if __name__ == '__main__':
    main()
