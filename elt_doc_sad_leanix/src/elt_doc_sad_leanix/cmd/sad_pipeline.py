#!/usr/bin/env python3
import sys
import argparse
from pathlib import Path
import xml.etree.ElementTree as ET
import openpyxl
from docx import Document
from elt_doc_sad_leanix.diagram_generator import WorkdayIntegrationDiagramGenerator


def extract_inventory(path: Path) -> str:
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if len(rows) < 2:
            return "Error: Inventory has no data rows"
        headers = [str(h) if h else "" for h in rows[0]]
        col_map = {}
        for i, h in enumerate(headers):
            h_lower = h.lower().strip()
            if h_lower == "name":
                col_map["name"] = i
            elif h_lower == "type":
                col_map["type"] = i
            elif h_lower == "id":
                col_map["id"] = i
        required = ["name", "type", "id"]
        missing = [c for c in required if c not in col_map]
        if missing:
            return f"Error: Inventory missing required columns: {missing}"
        lines = [
            "| name | type | id |",
            "|------|------|----|",
        ]
        for row in rows[2:]:
            name = row[col_map["name"]] or ""
            typ = row[col_map["type"]] or ""
            rid = row[col_map["id"]] or ""
            if name:
                lines.append(f"| {name} | {typ} | {rid} |")
        return "\n".join(lines)
    except Exception as e:
        return f"Error reading inventory: {e}"


def load_inventory_records(path: Path):
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if len(rows) < 2:
            return []
        headers = [str(h).strip().lower() if h else "" for h in rows[0]]
        name_idx = headers.index("name") if "name" in headers else -1
        type_idx = headers.index("type") if "type" in headers else -1
        id_idx = headers.index("id") if "id" in headers else -1
        if min(name_idx, type_idx, id_idx) < 0:
            return []
        recs = []
        for row in rows[2:]:
            n = (row[name_idx] or "").strip() if row[name_idx] else ""
            t = (row[type_idx] or "").strip() if row[type_idx] else ""
            i = (row[id_idx] or "").strip() if row[id_idx] else ""
            if n:
                recs.append({"name": n, "type": t, "id": i})
        return recs
    except Exception:
        return []


def lookup_record(records, name, type_filter=None):
    if not records or not name:
        return None, None
    nl = name.lower().strip()
    exact = [r for r in records if r["name"].lower() == nl]
    if type_filter:
        exact = [r for r in exact if r["type"] == type_filter]
    if exact:
        order = {"Application": 0, "Interface": 1, "ITComponent": 2, "Provider": 3}
        exact.sort(key=lambda r: order.get(r["type"], 99))
        r = exact[0]
        return r["id"], r["type"]
    contains = [r for r in records if nl in r["name"].lower()]
    if type_filter:
        contains = [r for r in contains if r["type"] == type_filter]
    if contains:
        order = {"Application": 0, "Interface": 1, "ITComponent": 2, "Provider": 3}
        contains.sort(key=lambda r: order.get(r["type"], 99))
        r = contains[0]
        return r["id"], r["type"]
    return None, None


def extract_sad_text(path: Path) -> str:
    try:
        doc = Document(path)
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text)
        for table in doc.tables:
            full_text.append("\n[TABLE START]")
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows.append(" | ".join(cells))
            full_text.append("\n".join(rows))
            full_text.append("[TABLE END]\n")
        return "\n".join(full_text)
    except Exception as e:
        return f"Error reading SAD: {e}"


def build_xml_from_template(spec: dict, template_path: Path) -> str:
    tree = ET.parse(template_path)
    root = tree.getroot()
    title = spec.get("title") or ""
    for cell in root.iter("mxCell"):
        if cell.get("vertex") == "1":
            style = cell.get("style") or ""
            if "text;strokeColor=none" in style and "fontSize=24" in style:
                if title:
                    cell.set("value", title)
                break
    objects = [o for o in root.iter("object") if o.get("type") == "factSheet"]
    source = None
    sftp = None
    target = None
    for o in objects:
        label = o.get("label") or ""
        if "Workday" in label and not source:
            source = o
        elif "SFTP" in label and not sftp:
            sftp = o
        else:
            if not target:
                target = o
    source_name = spec.get("source_system") or (source.get("label") if source is not None else "")
    target_name = spec.get("target_system") or (target.get("label") if target is not None else "")
    inter_name = spec.get("intermediary") or (sftp.get("label") if sftp is not None else "")
    integ_id = spec.get("integration_id") or ""
    if source is not None:
        sid = spec.get("source_id")
        if sid:
            source.set("factSheetId", sid)
        if source_name:
            source.set("label", source_name)
    if sftp is not None:
        iid = spec.get("intermediary_id")
        if iid:
            sftp.set("factSheetId", iid)
        if inter_name:
            base = inter_name
        else:
            base = "SFTP"
        if integ_id:
            sftp.set("label", f"{base}<div><b>{integ_id}</b></div>")
        else:
            sftp.set("label", base)
    if target is not None:
        tid = spec.get("target_id")
        if tid:
            target.set("factSheetId", tid)
        if target_name:
            target.set("label", target_name)
    processes = [
        spec.get("process_extraction") or "",
        spec.get("process_security") or "",
        spec.get("process_transmission") or "",
        spec.get("process_processing") or "",
    ]
    rows = [c for c in root.iter("mxCell") if c.get("style", "").startswith("shape=tableRow")]
    if len(rows) >= 2:
        body_row = rows[1]
        body_cells = [c for c in body_row if c.tag == "mxCell"]
        for idx, cell in enumerate(body_cells[:4]):
            text = processes[idx]
            if text:
                cell.set("value", text)
    return ET.tostring(root, encoding="unicode")


def build_prompt(sad_path: Path, prompt_path: Path, inventory_path: Path) -> str:
    if not prompt_path.exists():
        raise SystemExit(f"Error: Prompt file {prompt_path} not found")
    with open(prompt_path, "r", encoding="utf-8") as f:
        master_prompt = f.read()
    inventory_text = extract_inventory(inventory_path)
    sad_text = extract_sad_text(sad_path)
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent.parent.parent
    if not (project_root / "pyproject.toml").exists():
        project_root = script_dir.parent.parent.parent
    templates_dir = project_root / "elt_doc_sad_leanix" / "templates"
    if not templates_dir.exists():
        templates_dir = project_root / "templates"
    template_list = "No templates found."
    if templates_dir.exists():
        templates = [f.name for f in templates_dir.glob("*.xml")]
        template_list = "\n".join([f"* {t}" for t in sorted(templates)])
    full_output = f"""{master_prompt}

---

## CONTEXT DATA

### Available Templates
The following XML templates are available in the repository. You MUST select the one that best matches the integration pattern.

{template_list}

### LeanIX Inventory
The following is the authoritative list of systems and IDs. You MUST reference these IDs.

{inventory_text}

---

## INPUT DATA

### SAD Document Content
The following is the extracted text from the System Architecture Document.

{sad_text}
"""
    return full_output


def main():
    parser = argparse.ArgumentParser(description="SAD to LeanIX helper")
    parser.add_argument("sad_path", help="Path to SAD .docx file")
    parser.add_argument("--output-dir", help="Directory for prompt and XML outputs")
    parser.add_argument("--json-spec", help="Optional JSON spec path to build XML")
    args = parser.parse_args()

    sad_path = Path(args.sad_path)
    if not sad_path.exists():
        print(f"Error: SAD file {sad_path} not found")
        sys.exit(1)

    script_dir = Path(__file__).parent
    default_prompt = script_dir.parent / "prompts" / "sad_to_leanix.md"
    project_root = script_dir.parent.parent.parent.parent
    if not (project_root / "pyproject.toml").exists():
        project_root = script_dir.parent.parent.parent
    default_inventory = project_root / "config" / "LeanIX_Inventory.xlsx"
    if not default_inventory.exists():
        default_inventory = project_root / "elt_doc_sad_leanix" / "config" / "LeanIX_Inventory.xlsx"

    if args.output_dir:
        out_dir = Path(args.output_dir)
    else:
        out_dir = project_root / ".tmp"
    out_dir.mkdir(parents=True, exist_ok=True)
    prompts_dir = out_dir / "prompts"
    specs_dir = out_dir / "specs"
    xml_dir = out_dir / "xml"
    prompts_dir.mkdir(parents=True, exist_ok=True)
    specs_dir.mkdir(parents=True, exist_ok=True)
    xml_dir.mkdir(parents=True, exist_ok=True)

    prompt_text = build_prompt(sad_path, default_prompt, default_inventory)
    prompt_path = prompts_dir / f"{sad_path.stem}_prompt.md"
    prompt_path.write_text(prompt_text, encoding="utf-8")

    if not args.json_spec:
        abs_prompt = prompt_path.resolve()
        spec_path = (specs_dir / f"{sad_path.stem}_spec.json").resolve()
        print()
        print("Copy the following instruction into your LLM:")
        print("-" * 80)
        print(
            f"Process the prompt in `{abs_prompt}`, generate the JSON spec exactly as "
            f"described in the schema in that prompt, and save it to `{spec_path}`."
        )
        print("-" * 80)
        return

    json_path = Path(args.json_spec)
    if not json_path.is_absolute():
        if json_path.parent == Path("."):
            json_path = specs_dir / json_path.name
        else:
            json_path = out_dir / json_path
    if not json_path.exists():
        print(f"Error: JSON spec {json_path} not found")
        sys.exit(1)

    import json

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            spec = json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        sys.exit(1)

    required_fields = ["title", "integration_id", "source_system", "target_system"]
    missing = [f for f in required_fields if f not in spec]
    if missing:
        print(f"Error: Missing required fields in JSON: {missing}")
        sys.exit(1)

    inv_records = load_inventory_records(default_inventory)
    if inv_records:
        if not spec.get("source_id"):
            sid, stype = lookup_record(inv_records, spec.get("source_system"))
            if sid:
                spec["source_id"] = sid
                spec["source_type"] = stype
        if not spec.get("target_id"):
            tid, ttype = lookup_record(inv_records, spec.get("target_system"))
            if tid:
                spec["target_id"] = tid
                spec["target_type"] = ttype
        if spec.get("intermediary") and not spec.get("intermediary_id"):
            iid, itype = lookup_record(inv_records, spec.get("intermediary"))
            if iid:
                spec["intermediary_id"] = iid
                spec["intermediary_type"] = itype
        if not spec.get("interface_id"):
            label = spec.get("interface_label") or ""
            int_id = spec.get("integration_id") or ""
            cand = None
            for r in inv_records:
                if r["type"] == "Interface" and (int_id and int_id in r["name"] or label and label in r["name"]):
                    cand = r
                    break
            if cand:
                spec["interface_id"] = cand["id"]
                if not spec.get("interface_label"):
                    spec["interface_label"] = cand["name"]

    print(f"Generating XML for {spec.get('integration_id')} - {spec.get('title')}")
    template_id = spec.get("template_id")
    xml = None
    if template_id:
        print(f"Using Template Pattern: {template_id}")
        templates_dir = project_root / "elt_doc_sad_leanix" / "templates"
        if not templates_dir.exists():
            templates_dir = project_root / "templates"
        template_path = templates_dir / template_id
        if template_path.exists():
            xml = build_xml_from_template(spec, template_path)
    if xml is None:
        generator = WorkdayIntegrationDiagramGenerator()
        xml = generator.generate_xml(spec)

    output_path = xml_dir / f"{json_path.stem}.xml"
    output_path.write_text(xml, encoding="utf-8")
    print(f"Generated XML at {output_path}")


if __name__ == "__main__":
    main()
