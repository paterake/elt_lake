"""Content validator for SAD documents."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document

from .config_loader import load_sad_template, get_section_guidance

logger = logging.getLogger(__name__)

# Constants for validation
MIN_CONTENT_LENGTH = 100
HEADING_STYLE_PREFIX = "Heading"


class ValidationError(Exception):
    """Exception raised for validation errors."""
    pass


class DocumentLoadError(Exception):
    """Exception raised when document cannot be loaded."""
    pass


def validate_docx_path(doc_path: str) -> Path:
    """Validate and resolve document path.
    
    Args:
        doc_path: Path to the Word document.
        
    Returns:
        Resolved Path object.
        
    Raises:
        ValidationError: If path is invalid.
    """
    if not doc_path:
        raise ValidationError("Document path cannot be empty")
    
    try:
        path = Path(doc_path).expanduser().resolve()
        
        if not path.exists():
            raise ValidationError(f"Document not found: {path}")
        
        if not path.is_file():
            raise ValidationError(f"Path is not a file: {path}")
        
        if path.suffix.lower() != ".docx":
            raise ValidationError(f"File must be a .docx file: {path}")
        
        logger.debug("Validated document path: %s", path)
        return path
        
    except (OSError, RuntimeError) as e:
        raise ValidationError(f"Invalid document path '{doc_path}': {e}")


def extract_headings(doc_path: str) -> list[dict[str, Any]]:
    """Extract headings from a Word document.
    
    Args:
        doc_path: Path to the Word document.
        
    Returns:
        List of dictionaries containing heading information.
        
    Raises:
        DocumentLoadError: If document cannot be loaded.
    """
    try:
        logger.debug("Loading document: %s", doc_path)
        doc = Document(doc_path)
    except Exception as e:
        logger.error("Failed to load document: %s", e)
        raise DocumentLoadError(f"Failed to load document: {e}")
    
    headings = []
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        
        if style_name.startswith(HEADING_STYLE_PREFIX):
            try:
                level = int(style_name.split(' ')[-1]) if ' ' in style_name else 1
            except (ValueError, IndexError):
                level = 1
            
            headings.append({
                "level": level,
                "text": para.text.strip(),
                "style": style_name,
            })
    
    logger.debug("Extracted %d headings from document", len(headings))
    return headings


def validate_sad_structure(doc_path: str) -> dict[str, Any]:
    """Validate SAD document structure against template.
    
    Args:
        doc_path: Path to the SAD document.
        
    Returns:
        Validation result dictionary with status and issues.
        
    Raises:
        ValidationError: If document path is invalid.
        DocumentLoadError: If document cannot be loaded.
    """
    logger.info("Validating SAD structure: %s", doc_path)
    
    # Validate path first
    validated_path = validate_docx_path(doc_path)
    
    try:
        template = load_sad_template()
    except (FileNotFoundError, ValueError) as e:
        logger.error("Failed to load template: %s", e)
        raise ValidationError(f"Failed to load template: {e}")
    
    try:
        headings = extract_headings(str(validated_path))
    except DocumentLoadError as e:
        logger.error("Failed to extract headings: %s", e)
        raise
    
    heading_texts = [h["text"] for h in headings]
    
    issues = {
        "missing_required": [],
        "missing_optional": [],
        "found": [],
    }
    
    # Check front matter
    front_matter = template.get("front_matter", [])
    for item in front_matter:
        name = item.get("name", "")
        required = item.get("required", False)
        
        if name in heading_texts:
            issues["found"].append(f"Front Matter: {name}")
        elif required:
            issues["missing_required"].append(f"Front Matter: {name}")
        else:
            issues["missing_optional"].append(f"Front Matter: {name}")
    
    # Check sections
    sections = template.get("sections", {})
    for section_num, section_data in sections.items():
        section_title = section_data.get("title", "")
        full_title = f"{section_num}. {section_title}"
        
        # Check main section
        if any(full_title in h for h in heading_texts):
            issues["found"].append(full_title)
        else:
            issues["missing_required"].append(full_title)
        
        # Check subsections
        subsections = section_data.get("subsections", {})
        for sub_id, sub_data in subsections.items():
            sub_title = sub_data.get("title", "")
            required = sub_data.get("required", False)
            
            # Format subsection number (e.g., "1.1", "2.3")
            if '.' in sub_id:
                full_sub_id = f"{section_num}.{sub_id.split('.')[1]}"
            else:
                full_sub_id = f"{section_num}.{sub_id}"
            
            full_sub_title = f"{full_sub_id} {sub_title}"
            
            if any(full_sub_title in h for h in heading_texts):
                issues["found"].append(full_sub_title)
            elif required:
                issues["missing_required"].append(full_sub_title)
            else:
                issues["missing_optional"].append(full_sub_title)
    
    result = {
        "valid": len(issues["missing_required"]) == 0,
        "issues": issues,
        "completeness": calculate_completeness(issues, template),
        "total_headings": len(headings),
    }
    
    logger.info(
        "Validation complete: valid=%s, completeness=%.1f%%",
        result["valid"],
        result["completeness"],
    )
    
    return result


def calculate_completeness(issues: dict[str, Any], template: dict[str, Any]) -> float:
    """Calculate document completeness percentage.
    
    Args:
        issues: Issues dictionary from validation.
        template: SAD template structure.
        
    Returns:
        Completeness percentage (0-100).
    """
    total_required = 0
    found_count = len(issues["found"])
    
    # Count front matter
    for item in template.get("front_matter", []):
        if item.get("required", False):
            total_required += 1
    
    # Count sections
    for section_data in template.get("sections", {}).values():
        total_required += 1  # Main section
        for sub_data in section_data.get("subsections", {}).values():
            if sub_data.get("required", False):
                total_required += 1
    
    if total_required == 0:
        return 100.0
    
    return round((found_count / total_required) * 100, 1)


def validate_section_content(doc_path: str, section_id: str) -> dict[str, Any]:
    """Validate content of a specific section.
    
    Args:
        doc_path: Path to the SAD document.
        section_id: Section identifier (e.g., "1.1", "2.3").
        
    Returns:
        Validation result with content assessment.
        
    Raises:
        ValidationError: If document path is invalid or section not found.
        DocumentLoadError: If document cannot be loaded.
    """
    logger.debug("Validating section content: %s", section_id)
    
    validated_path = validate_docx_path(doc_path)
    
    try:
        guidance = get_section_guidance(section_id)
    except (FileNotFoundError, ValueError) as e:
        logger.error("Failed to load section guidance: %s", e)
        raise ValidationError(f"Failed to load section guidance: {e}")
    
    if not guidance:
        return {
            "valid": False,
            "error": f"No guidance found for section {section_id}",
        }
    
    try:
        doc = Document(str(validated_path))
    except Exception as e:
        logger.error("Failed to load document: %s", e)
        raise DocumentLoadError(f"Failed to load document: {e}")
    
    # Find section in document
    in_section = False
    section_content = []
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        
        if style_name.startswith(HEADING_STYLE_PREFIX):
            if section_id in para.text:
                in_section = True
                continue
            elif in_section:
                # Check if this is a subsection or next section
                try:
                    level = int(style_name.split(' ')[-1]) if ' ' in style_name else 1
                    if level <= 2:
                        break
                except (ValueError, IndexError):
                    pass
        elif in_section and para.text.strip():
            section_content.append(para.text.strip())
    
    content_length = sum(len(c) for c in section_content)
    should_include = guidance.get("should_include", [])
    
    result = {
        "valid": content_length > MIN_CONTENT_LENGTH,
        "content_length": content_length,
        "paragraphs": len(section_content),
        "should_include_count": len(should_include),
        "guidance": guidance.get("description", ""),
        "min_length": MIN_CONTENT_LENGTH,
    }
    
    logger.debug(
        "Section %s validation: valid=%s, length=%d",
        section_id,
        result["valid"],
        content_length,
    )
    
    return result


def generate_validation_report(doc_path: str) -> str:
    """Generate a human-readable validation report.
    
    Args:
        doc_path: Path to the SAD document.
        
    Returns:
        Formatted validation report string.
        
    Raises:
        ValidationError: If document path is invalid.
        DocumentLoadError: If document cannot be loaded.
    """
    result = validate_sad_structure(doc_path)
    
    lines = [
        "=" * 60,
        "SAD Document Validation Report",
        "=" * 60,
        f"Document: {doc_path}",
        f"Valid: {'Yes' if result['valid'] else 'No'}",
        f"Completeness: {result['completeness']}%",
        f"Total Headings: {result.get('total_headings', 0)}",
        "",
    ]
    
    if result["issues"]["found"]:
        lines.append(f"Found Sections ({len(result['issues']['found'])}):")
        for item in result["issues"]["found"]:
            lines.append(f"  ✓ {item}")
        lines.append("")
    
    if result["issues"]["missing_required"]:
        lines.append(
            f"Missing Required Sections ({len(result['issues']['missing_required'])}):"
        )
        for item in result["issues"]["missing_required"]:
            lines.append(f"  ✗ {item}")
        lines.append("")
    
    if result["issues"]["missing_optional"]:
        lines.append(
            f"Missing Optional Sections ({len(result['issues']['missing_optional'])}):"
        )
        for item in result["issues"]["missing_optional"]:
            lines.append(f"  ○ {item}")
        lines.append("")
    
    lines.append("=" * 60)
    
    return "\n".join(lines)


def main() -> None:
    """CLI entry point for SAD validation."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Validate SAD document")
    parser.add_argument(
        "--input",
        required=True,
        help="Input SAD document path",
    )
    parser.add_argument(
        "--section",
        help="Validate specific section only (e.g., 1.1)",
    )
    parser.add_argument(
        "--report",
        action="store_true",
        help="Generate full report",
    )
    parser.add_argument(
        "--verbose",
        "-v",
        action="store_true",
        help="Enable verbose logging",
    )
    
    args = parser.parse_args()
    
    # Configure logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )
    
    try:
        if args.section:
            result = validate_section_content(args.input, args.section)
            print(f"Section {args.section} Validation:")
            for key, value in result.items():
                print(f"  {key}: {value}")
        elif args.report:
            report = generate_validation_report(args.input)
            print(report)
        else:
            result = validate_sad_structure(args.input)
            print(f"Valid: {result['valid']}")
            print(f"Completeness: {result['completeness']}%")
            if result["issues"]["missing_required"]:
                print("\nMissing required sections:")
                for item in result["issues"]["missing_required"]:
                    print(f"  - {item}")
                    
    except ValidationError as e:
        print(f"Validation error: {e}")
        raise SystemExit(1)
    except DocumentLoadError as e:
        print(f"Document load error: {e}")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
