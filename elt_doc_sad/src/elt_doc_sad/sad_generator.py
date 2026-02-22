"""SAD document generator."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config_loader import (
    load_sad_template,
    load_section_guidance,
    get_cover_image_path,
)

logger = logging.getLogger(__name__)

# Constants for document formatting
TITLE_FONT_SIZE = Pt(24)
SUBTITLE_FONT_SIZE = Pt(16)
COVER_IMAGE_WIDTH = Inches(3)
DEFAULT_VERSION = "1.0"

# Validation patterns
INTEGRATION_ID_PATTERN = re.compile(r"^[A-Z]{2,5}\d{3,5}$")
VENDOR_NAME_MIN_LENGTH = 2
VENDOR_NAME_MAX_LENGTH = 100


class SadGenerationError(Exception):
    """Exception raised for SAD generation errors."""
    pass


class ValidationError(Exception):
    """Exception raised for input validation errors."""
    pass


def _get_package_root() -> Path:
    """Get the package root directory path."""
    return Path(__file__).resolve().parent


def validate_integration_id(integration_id: str) -> str:
    """Validate and normalize integration ID.
    
    Args:
        integration_id: Integration identifier to validate.
        
    Returns:
        Normalized integration ID.
        
    Raises:
        ValidationError: If integration ID is invalid.
    """
    if not integration_id:
        raise ValidationError("Integration ID cannot be empty")
    
    normalized = integration_id.strip().upper()
    
    if not INTEGRATION_ID_PATTERN.match(normalized):
        raise ValidationError(
            f"Invalid integration ID format: '{integration_id}'. "
            f"Expected format: 2-5 uppercase letters followed by 3-5 digits (e.g., INT001)"
        )
    
    logger.debug("Validated integration ID: %s", normalized)
    return normalized


def validate_vendor_name(vendor_name: str) -> str:
    """Validate vendor name.
    
    Args:
        vendor_name: Vendor name to validate.
        
    Returns:
        Validated vendor name.
        
    Raises:
        ValidationError: If vendor name is invalid.
    """
    if not vendor_name:
        raise ValidationError("Vendor name cannot be empty")
    
    normalized = vendor_name.strip()
    
    if len(normalized) < VENDOR_NAME_MIN_LENGTH:
        raise ValidationError(
            f"Vendor name too short: '{vendor_name}'. "
            f"Minimum length is {VENDOR_NAME_MIN_LENGTH} characters"
        )
    
    if len(normalized) > VENDOR_NAME_MAX_LENGTH:
        raise ValidationError(
            f"Vendor name too long: '{vendor_name}'. "
            f"Maximum length is {VENDOR_NAME_MAX_LENGTH} characters"
        )
    
    logger.debug("Validated vendor name: %s", normalized)
    return normalized


def validate_output_path(output_path: str) -> Path:
    """Validate and resolve output path.
    
    Args:
        output_path: Output path string.
        
    Returns:
        Resolved Path object.
        
    Raises:
        ValidationError: If output path is invalid.
    """
    if not output_path:
        raise ValidationError("Output path cannot be empty")
    
    try:
        expanded = Path(output_path).expanduser()
        resolved = expanded.resolve()
        logger.debug("Resolved output path: %s", resolved)
        return resolved
    except (OSError, RuntimeError) as e:
        raise ValidationError(f"Invalid output path '{output_path}': {e}")


def create_cover_page(
    doc: Document,
    title: str,
    integration_id: str,
    vendor_name: str,
) -> None:
    """Create a cover page for the SAD document.
    
    Args:
        doc: Word document to add cover page to.
        title: Document title.
        integration_id: Integration identifier (e.g., INT001).
        vendor_name: Vendor company name.
    """
    logger.debug("Creating cover page with title: %s", title)
    
    # Add title
    heading = doc.add_heading(level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.font.size = TITLE_FONT_SIZE
    run.font.bold = True
    
    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Solution Architecture Definition")
    subtitle_run.font.size = SUBTITLE_FONT_SIZE
    
    # Add integration details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run(f"\nIntegration ID: {integration_id}\n")
    details.add_run(f"Vendor: {vendor_name}\n")
    details.add_run(f"Version: {DEFAULT_VERSION}\n")
    
    # Add cover image if available
    try:
        cover_image = get_cover_image_path()
        if cover_image.exists():
            logger.debug("Adding cover image from: %s", cover_image)
            doc.add_picture(str(cover_image), width=COVER_IMAGE_WIDTH)
            # Center the image
            for paragraph in doc.paragraphs:
                if paragraph.runs and paragraph.runs[-1].text == '':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logger.warning("Cover image not found: %s", cover_image)
    except Exception as e:
        logger.warning("Failed to add cover image: %s", e)


def create_document_history(doc: Document) -> None:
    """Create the document history table.
    
    Args:
        doc: Word document to add table to.
    """
    logger.debug("Creating document history table")
    
    doc.add_heading("Document History", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ["Version", "Date", "Author", "Changes"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # First data row (initial version)
    row = table.rows[1]
    row.cells[0].text = DEFAULT_VERSION
    row.cells[1].text = "[Date]"
    row.cells[2].text = "[Author]"
    row.cells[3].text = "Initial version"


def create_document_review(doc: Document) -> None:
    """Create the document review table.
    
    Args:
        doc: Word document to add table to.
    """
    logger.debug("Creating document review table")
    
    doc.add_heading("Document Review", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ["Name", "Role", "Date", "Signature"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True


def create_approvals(doc: Document) -> None:
    """Create the approvals table.
    
    Args:
        doc: Word document to add table to.
    """
    logger.debug("Creating approvals table")
    
    doc.add_heading("Approvals", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ["Name", "Role", "Date", "Signature"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True


def create_section_placeholder(
    doc: Document,
    section_num: str,
    section_id: str,
    section_data: dict[str, Any],
    guidance_cache: dict[str, Any] | None = None,
) -> None:
    """Create a section with placeholder content.
    
    Args:
        doc: Word document to add section to.
        section_num: Section number (e.g., "1", "2").
        section_id: Subsection ID (e.g., "1", "1.1", "2.3").
        section_data: Section configuration data.
        guidance_cache: Optional cache of section guidance.
    """
    title = section_data.get("title", f"Section {section_id}")
    level = int(section_data.get("level", 2))
    
    # Create full heading with number
    if section_id == section_num:
        heading_text = f"{section_num}. {title}"
    else:
        heading_text = f"{section_id} {title}"
    
    logger.debug("Creating section: %s", heading_text)
    heading = doc.add_heading(heading_text, level=level)
    
    # Use cached guidance or load fresh
    if guidance_cache is None:
        try:
            guidance = load_section_guidance()
        except (FileNotFoundError, ValueError):
            logger.warning("Could not load section guidance")
            return
    else:
        guidance = guidance_cache
    
    # Find matching guidance
    guidance_key = None
    for key in guidance:
        if key.startswith(section_id):
            guidance_key = key
            break
    
    if guidance_key and guidance_key in guidance:
        guidance_data = guidance[guidance_key]
        description = guidance_data.get("description", "")
        should_include = guidance_data.get("should_include", [])
        
        if description:
            doc.add_paragraph(f"[{description}]")
        
        if should_include:
            doc.add_paragraph("This section should include:")
            for item in should_include:
                doc.add_paragraph(f"  • {item}", style="List Bullet")


def generate_sad_document(
    output_path: str,
    integration_id: str,
    vendor_name: str,
    title: str | None = None,
) -> Path:
    """Generate a complete SAD document.
    
    Args:
        output_path: Path to save the document (directory or full file path).
        integration_id: Integration identifier (e.g., INT001).
        vendor_name: Vendor company name.
        title: Optional custom title. Defaults to standard SAD title.
        
    Returns:
        Path to the generated document.
        
    Raises:
        ValidationError: If input parameters are invalid.
        SadGenerationError: If document generation fails.
    """
    logger.info(
        "Generating SAD document: integration_id=%s, vendor_name=%s",
        integration_id,
        vendor_name,
    )
    
    # Validate inputs
    try:
        validated_id = validate_integration_id(integration_id)
        validated_vendor = validate_vendor_name(vendor_name)
        validated_output = validate_output_path(output_path)
    except ValidationError as e:
        logger.error("Input validation failed: %s", e)
        raise
    
    # Set default title
    if title is None:
        title = f"SAD_{validated_id}_{validated_vendor.replace(' ', '_')}_V{DEFAULT_VERSION.replace('.', '_')}"
        logger.debug("Generated default title: %s", title)
    
    try:
        # Create document
        logger.debug("Creating new Word document")
        doc = Document()
        
        # Create cover page
        create_cover_page(doc, title, validated_id, validated_vendor)
        
        # Add page break
        doc.add_page_break()
        
        # Front matter
        logger.debug("Adding front matter")
        create_document_history(doc)
        doc.add_paragraph()
        create_document_review(doc)
        doc.add_paragraph()
        create_approvals(doc)
        
        # Load template structure
        logger.debug("Loading SAD template")
        template = load_sad_template()
        
        # Cache guidance for performance
        try:
            guidance_cache = load_section_guidance()
        except (FileNotFoundError, ValueError) as e:
            logger.warning("Could not load section guidance: %s", e)
            guidance_cache = {}
        
        # Generate sections
        section_count = 0
        subsection_count = 0
        
        for section_num, section_data in template.get("sections", {}).items():
            doc.add_page_break()
            
            # Main section heading
            section_title = section_data.get("title", f"Section {section_num}")
            doc.add_heading(f"{section_num}. {section_title}", level=1)
            section_count += 1
            
            # Subsections
            subsections = section_data.get("subsections", {})
            for sub_id, sub_data in subsections.items():
                full_id = f"{section_num}.{sub_id.split('.')[1]}" if '.' in sub_id else sub_id
                create_section_placeholder(
                    doc, section_num, full_id, sub_data, guidance_cache
                )
                subsection_count += 1
        
        logger.info(
            "Generated document with %d sections and %d subsections",
            section_count,
            subsection_count,
        )
        
        # Determine output path
        output = validated_output
        if output.is_dir():
            output = output / f"{title}.docx"
        elif not output.suffix:
            output = output / f"{title}.docx"
        
        # Ensure parent directory exists
        try:
            output.parent.mkdir(parents=True, exist_ok=True)
            logger.debug("Created output directory: %s", output.parent)
        except OSError as e:
            raise SadGenerationError(
                f"Failed to create output directory {output.parent}: {e}"
            )
        
        # Save document
        logger.info("Saving document to: %s", output)
        doc.save(str(output))
        
        # Verify file was created
        if not output.exists():
            raise SadGenerationError(f"Document was not saved to {output}")
        
        logger.info("SAD document generated successfully: %s", output)
        return output
        
    except SadGenerationError:
        raise
    except Exception as e:
        logger.exception("Unexpected error during document generation")
        raise SadGenerationError(f"Failed to generate SAD document: {e}")


def main() -> None:
    """CLI entry point for SAD generation."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Generate SAD document")
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Output directory",
    )
    parser.add_argument(
        "--integration-id",
        required=True,
        help="Integration ID (e.g., INT001)",
    )
    parser.add_argument(
        "--vendor-name",
        required=True,
        help="Vendor name",
    )
    parser.add_argument(
        "--title",
        help="Custom title (optional)",
    )
    parser.add_argument(
        "--verbose",
        "-v",
        action="store_true",
        help="Enable verbose logging",
    )
    
    args = parser.parse_args()
    
    # Configure logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )
    
    try:
        output_path = generate_sad_document(
            output_path=args.output_dir,
            integration_id=args.integration_id,
            vendor_name=args.vendor_name,
            title=args.title,
        )
        print(f"SAD document generated: {output_path}")
    except ValidationError as e:
        print(f"Validation error: {e}")
        raise SystemExit(1)
    except SadGenerationError as e:
        print(f"Generation error: {e}")
        raise SystemExit(1)


if __name__ == "__main__":
    main()
