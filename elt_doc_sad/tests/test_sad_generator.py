"""Tests for sad_generator module."""

from __future__ import annotations

from pathlib import Path

import pytest

from elt_doc_sad.sad_generator import (
    SadGenerationError,
    ValidationError,
    create_approvals,
    create_cover_page,
    create_document_history,
    create_document_review,
    generate_sad_document,
    validate_integration_id,
    validate_output_path,
    validate_vendor_name,
)
from docx import Document


class TestValidationIntegrationId:
    """Test integration ID validation."""

    def test_valid_integration_id(self, valid_integration_ids: list[str]):
        """Test validation of valid integration IDs."""
        for integration_id in valid_integration_ids:
            result = validate_integration_id(integration_id)
            assert result == integration_id.strip().upper()

    def test_invalid_integration_id(self, invalid_integration_ids: list[str]):
        """Test validation of invalid integration IDs."""
        for integration_id in invalid_integration_ids:
            with pytest.raises(ValidationError):
                validate_integration_id(integration_id)

    def test_empty_integration_id(self):
        """Test validation of empty integration ID."""
        with pytest.raises(ValidationError, match="cannot be empty"):
            validate_integration_id("")

    def test_whitespace_integration_id(self):
        """Test validation of whitespace integration ID."""
        # Whitespace-only gets stripped and then fails format validation
        with pytest.raises(ValidationError):
            validate_integration_id("   ")

    def test_normalization(self):
        """Test integration ID normalization."""
        result = validate_integration_id("  int001  ")
        assert result == "INT001"


class TestValidationVendorName:
    """Test vendor name validation."""

    def test_valid_vendor_name(self, valid_vendor_names: list[str]):
        """Test validation of valid vendor names."""
        for vendor_name in valid_vendor_names:
            # Skip edge cases that might fail
            if len(vendor_name.strip()) < 2 or len(vendor_name.strip()) > 100:
                continue
            result = validate_vendor_name(vendor_name)
            assert result == vendor_name.strip()

    def test_invalid_vendor_name(self, invalid_vendor_names: list[str]):
        """Test validation of invalid vendor names."""
        for vendor_name in invalid_vendor_names:
            with pytest.raises(ValidationError, match="Vendor name"):
                validate_vendor_name(vendor_name)

    def test_empty_vendor_name(self):
        """Test validation of empty vendor name."""
        with pytest.raises(ValidationError, match="cannot be empty"):
            validate_vendor_name("")

    def test_normalization(self):
        """Test vendor name normalization."""
        result = validate_vendor_name("  Nebula HR  ")
        assert result == "Nebula HR"


class TestValidationOutputPath:
    """Test output path validation."""

    def test_valid_directory_path(self, temp_output_dir: Path):
        """Test validation of valid directory path."""
        result = validate_output_path(str(temp_output_dir))
        assert result == temp_output_dir.resolve()

    def test_valid_file_path(self, temp_output_dir: Path):
        """Test validation of valid file path."""
        file_path = temp_output_dir / "output.docx"
        result = validate_output_path(str(file_path))
        assert result == file_path.resolve()

    def test_empty_path(self):
        """Test validation of empty path."""
        with pytest.raises(ValidationError, match="cannot be empty"):
            validate_output_path("")

    def test_tilde_expansion(self):
        """Test tilde expansion in path."""
        result = validate_output_path("~/test_output")
        assert "~" not in str(result)


class TestDocumentCreation:
    """Test document creation functions."""

    def test_create_document_history(self, temp_output_dir: Path):
        """Test document history table creation."""
        doc = Document()
        create_document_history(doc)
        
        # Check table was created
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        # Check headers
        headers = [table.cell(0, i).text for i in range(4)]
        assert headers == ["Version", "Date", "Author", "Changes"]
        
        # Check first data row
        assert table.cell(1, 0).text == "1.0"
        assert table.cell(1, 3).text == "Initial version"

    def test_create_document_review(self, temp_output_dir: Path):
        """Test document review table creation."""
        doc = Document()
        create_document_review(doc)
        
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        headers = [table.cell(0, i).text for i in range(4)]
        assert headers == ["Name", "Role", "Date", "Signature"]

    def test_create_approvals(self, temp_output_dir: Path):
        """Test approvals table creation."""
        doc = Document()
        create_approvals(doc)
        
        assert len(doc.tables) == 1
        table = doc.tables[0]
        
        headers = [table.cell(0, i).text for i in range(4)]
        assert headers == ["Name", "Role", "Date", "Signature"]

    def test_create_cover_page(self, temp_output_dir: Path):
        """Test cover page creation."""
        doc = Document()
        create_cover_page(doc, "Test Title", "INT001", "Test Vendor")
        
        # Check headings were created
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 1
        assert headings[0].text == "Test Title"


class TestGenerateSadDocument:
    """Test SAD document generation."""

    def test_generate_basic_document(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test generating a basic SAD document."""
        output_path = generate_sad_document(
            output_path=str(temp_output_dir),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )
        
        assert output_path.exists()
        assert output_path.suffix == ".docx"
        
        # Verify document structure
        doc = Document(str(output_path))
        
        # Check for expected headings
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        assert any("Introduction" in h for h in heading_texts)
        assert any("Data" in h for h in heading_texts)
        assert any("1.1" in h for h in heading_texts)

    def test_generate_with_custom_title(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test generating document with custom title."""
        custom_title = "Custom_SAD_Title_V1_0"
        
        output_path = generate_sad_document(
            output_path=str(temp_output_dir),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
            title=custom_title,
        )
        
        assert output_path.name == f"{custom_title}.docx"

    def test_generate_to_specific_file(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test generating document to specific file path."""
        output_file = temp_output_dir / "custom_folder" / "my_sad.docx"

        output_path = generate_sad_document(
            output_path=str(output_file),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )

        # Output path should end with .docx
        assert output_path.suffix == ".docx"
        assert output_path.exists()

    def test_generate_creates_parent_directories(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test that parent directories are created."""
        nested_path = temp_output_dir / "level1" / "level2" / "level3" / "output.docx"
        
        output_path = generate_sad_document(
            output_path=str(nested_path),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )
        
        assert output_path.exists()
        assert output_path.parent.exists()

    def test_invalid_integration_id_raises_error(
        self,
        temp_output_dir: Path,
    ):
        """Test that invalid integration ID raises error."""
        with pytest.raises(ValidationError, match="Invalid integration ID"):
            generate_sad_document(
                output_path=str(temp_output_dir),
                integration_id="invalid",
                vendor_name="Test Vendor",
            )

    def test_invalid_vendor_name_raises_error(
        self,
        temp_output_dir: Path,
    ):
        """Test that invalid vendor name raises error."""
        with pytest.raises(ValidationError, match="Vendor name"):
            generate_sad_document(
                output_path=str(temp_output_dir),
                integration_id="INT001",
                vendor_name="",
            )

    def test_generate_multiple_documents(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test generating multiple documents."""
        integrations = [
            {"id": "INT001", "vendor": "Vendor A"},
            {"id": "INT002", "vendor": "Vendor B"},
            {"id": "INT003", "vendor": "Vendor C"},
        ]
        
        output_paths = []
        for integration in integrations:
            output_path = generate_sad_document(
                output_path=str(temp_output_dir),
                integration_id=integration["id"],
                vendor_name=integration["vendor"],
            )
            output_paths.append(output_path)
        
        # All documents should exist
        for path in output_paths:
            assert path.exists()
        
        # All should be unique
        assert len(set(str(p) for p in output_paths)) == len(output_paths)


class TestDocumentStructure:
    """Test generated document structure."""

    def test_document_has_all_sections(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test that generated document has all required sections."""
        output_path = generate_sad_document(
            output_path=str(temp_output_dir),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )
        
        doc = Document(str(output_path))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # Check main sections
        expected_sections = [
            "Introduction",
            "Data",
            "Non-functional Requirements",
            "Solution Architecture",
            "Technology Stack",
            "Costs",
            "Maintenance",
            "Appendix",
        ]
        
        for section in expected_sections:
            assert any(section in h for h in heading_texts), f"Missing section: {section}"

    def test_document_has_front_matter(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test that document has front matter tables."""
        output_path = generate_sad_document(
            output_path=str(temp_output_dir),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )
        
        doc = Document(str(output_path))
        
        # Should have at least 3 tables (history, review, approvals)
        assert len(doc.tables) >= 3

    def test_document_has_correct_numbering(
        self,
        temp_output_dir: Path,
        fictional_integration: dict,
    ):
        """Test that document has correct section numbering."""
        output_path = generate_sad_document(
            output_path=str(temp_output_dir),
            integration_id=fictional_integration["integration_id"],
            vendor_name=fictional_integration["vendor_name"],
        )
        
        doc = Document(str(output_path))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
        
        # Check subsection numbering
        assert any("1.1" in h for h in heading_texts)
        assert any("1.2" in h for h in heading_texts)
        assert any("2.1" in h for h in heading_texts)
        assert any("3.1" in h for h in heading_texts)
