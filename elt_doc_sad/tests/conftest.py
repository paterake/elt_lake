"""Pytest fixtures and configuration for elt_doc_sad tests.

This module provides shared fixtures for testing the SAD generation module.
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from typing import Any

import pytest


# =============================================================================
# Fictional Test Scenario: Nebula HR Integration
# =============================================================================
# This is a completely fabricated integration scenario for testing purposes.
# It does not reference any real Workday integration projects.
# =============================================================================

FICTIONAL_INTEGRATION_SPEC = {
    "integration_id": "NEB001",
    "vendor_name": "Nebula HR Solutions",
    "pattern": "outbound_eib_sftp",
    "pattern_name": "Outbound EIB to SFTP",
    "data_flow": "Workday → [EIB Custom Report] → [CSV] → [PGP Encrypt] → [SFTP] → Nebula",
    "description": "Employee data provisioning to Nebula HR platform",
    "data_elements": [
        "Employee ID",
        "Legal Name",
        "Email Address",
        "Job Title",
        "Department",
        "Manager",
        "Start Date",
        "Employment Type",
    ],
    "security_requirements": [
        "PGP encryption using Nebula public key",
        "SFTP with SSH key authentication",
        "TLS 1.2+ for data in transit",
        "Data minimization - only required fields",
    ],
    "components": [
        {"name": "EIB", "role": "Extract"},
        {"name": "Document Transformation", "role": "Transform"},
        {"name": "PGP Encryption", "role": "Secure"},
        {"name": "SFTP", "role": "Deliver"},
    ],
}


@pytest.fixture
def fictional_integration() -> dict[str, Any]:
    """Return the fictional Nebula HR integration specification.
    
    This fixture provides a realistic but completely fabricated integration
    specification for testing SAD generation and validation.
    
    Returns:
        Dictionary containing integration specification.
    """
    return FICTIONAL_INTEGRATION_SPEC.copy()


@pytest.fixture
def temp_output_dir() -> Path:
    """Create a temporary directory for test outputs.
    
    Yields:
        Path to temporary directory.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_docx_path(temp_output_dir: Path) -> Path:
    """Create a sample Word document for testing.
    
    Args:
        temp_output_dir: Temporary directory fixture.
        
    Yields:
        Path to sample document.
    """
    from docx import Document
    
    doc = Document()
    doc.add_heading("Test SAD Document", level=1)
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Objectives", level=2)
    doc.add_paragraph("Test objectives content for validation.")
    doc.add_heading("1.2 Functionality", level=2)
    doc.add_paragraph("Test functionality content.")
    
    output_path = temp_output_dir / "test_sample.docx"
    doc.save(str(output_path))
    
    yield output_path


@pytest.fixture
def config_dir() -> Path:
    """Get the config directory path.
    
    Returns:
        Path to config directory.
    """
    return Path(__file__).resolve().parent.parent / "config"


@pytest.fixture
def mock_sad_template() -> dict[str, Any]:
    """Return a mock SAD template for testing.
    
    Returns:
        Mock template dictionary.
    """
    return {
        "version": "2.0",
        "front_matter": [
            {"name": "Document History", "type": "table", "required": True},
            {"name": "Document Review", "type": "table", "required": True},
            {"name": "Approvals", "type": "table", "required": True},
            {"name": "Reference Documents", "type": "list", "required": True},
            {"name": "Guidance", "type": "text", "required": False},
        ],
        "sections": {
            "1": {
                "title": "Introduction",
                "level": 1,
                "subsections": {
                    "1.1": {"title": "Objectives", "required": True, "level": 2},
                    "1.2": {"title": "Functionality", "required": True, "level": 2},
                    "1.3": {"title": "Constraints", "required": True, "level": 2},
                    "1.4": {"title": "Dependencies", "required": True, "level": 2},
                    "1.5": {"title": "Legacy", "required": False, "level": 2},
                },
            },
            "2": {
                "title": "Data",
                "level": 1,
                "subsections": {
                    "2.1": {"title": "GDPR", "required": True, "level": 2},
                    "2.2": {"title": "Sources", "required": True, "level": 2},
                    "2.3": {"title": "Integration", "required": True, "level": 2},
                },
            },
        },
    }


@pytest.fixture
def mock_section_guidance() -> dict[str, Any]:
    """Return mock section guidance for testing.
    
    Returns:
        Mock guidance dictionary.
    """
    return {
        "1.1 Objectives": {
            "description": "Primary objectives of the integration",
            "should_include": [
                "Business goals",
                "Technical objectives",
                "Success criteria",
            ],
            "example_topics": [
                "Automated data provisioning",
                "Secure data transmission",
            ],
            "typical_length": "5-10 bullet points",
        },
        "1.2 Functionality": {
            "description": "Detailed description of solution functionality",
            "should_include": [
                "Component breakdown",
                "Data flow description",
            ],
            "example_topics": [
                "EIB custom report",
                "File format and encoding",
            ],
        },
    }


@pytest.fixture
def temp_config_dir(
    mock_sad_template: dict[str, Any],
    mock_section_guidance: dict[str, Any],
) -> Path:
    """Create a temporary directory with mock config files.
    
    Args:
        mock_sad_template: Mock template fixture.
        mock_section_guidance: Mock guidance fixture.
        
    Yields:
        Path to temporary config directory.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        config_path = Path(tmpdir)
        
        # Write mock template
        template_file = config_path / "sad_template.json"
        with template_file.open("w") as f:
            json.dump(mock_sad_template, f, indent=2)
        
        # Write mock guidance
        guidance_file = config_path / "section_guidance.json"
        with guidance_file.open("w") as f:
            json.dump(mock_section_guidance, f, indent=2)
        
        # Write mock patterns
        patterns = {
            "patterns": {
                "outbound_eib_sftp": {
                    "name": "Outbound EIB to SFTP",
                    "description": "Workday → EIB → SFTP → Vendor",
                },
            }
        }
        patterns_file = config_path / "integration_patterns.json"
        with patterns_file.open("w") as f:
            json.dump(patterns, f, indent=2)
        
        yield config_path


# =============================================================================
# Test Data Files
# =============================================================================

@pytest.fixture
def valid_integration_ids() -> list[str]:
    """Return list of valid integration ID formats.
    
    Returns:
        List of valid integration IDs.
    """
    return ["INT001", "NEB001", "WORK123", "ABC9999", "XY12345"]


@pytest.fixture
def invalid_integration_ids() -> list[str]:
    """Return list of invalid integration ID formats.
    
    Returns:
        List of invalid integration IDs.
    """
    return [
        "",
        "   ",
        "123",
        "int001",  # lowercase
        "INT1",  # too few digits
        "INT00001",  # too many digits
        "1INT001",  # starts with number
        "INT-001",  # contains hyphen
        "Integration_001",  # too long
    ]


@pytest.fixture
def valid_vendor_names() -> list[str]:
    """Return list of valid vendor names.
    
    Returns:
        List of valid vendor names.
    """
    return [
        "Nebula HR",
        "Okta",
        "Workday",
        "SAP SuccessFactors",
        "A",  # Minimum length
        "A" * 100,  # Maximum length
    ]


@pytest.fixture
def invalid_vendor_names() -> list[str]:
    """Return list of invalid vendor names.
    
    Returns:
        List of invalid vendor names.
    """
    return [
        "",
        "   ",
        "A" * 101,  # Too long
    ]
