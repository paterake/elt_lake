"""Report generator - Word document output."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..models.assessment import (
    AssessmentConfig,
    AssessmentResult,
    Finding,
    Recommendation,
    SectionResult,
    Severity,
    Status,
)


class ReportGenerator:
    """Generate Word document assessment report."""

    def __init__(self, config: AssessmentConfig):
        self.config = config
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Set up document styles."""
        # Title style
        styles = self.doc.styles
        if "AssessmentTitle" not in styles:
            title_style = styles.add_style("AssessmentTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = "Calibri"
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)

        # Heading style
        if "AssessmentHeading" not in styles:
            heading_style = styles.add_style("AssessmentHeading", WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = "Calibri"
            heading_font.size = Pt(16)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)

    def generate(self, results: list[AssessmentResult], screenshots_dir: Path) -> Path:
        """Generate the full assessment report."""
        self._add_title_page()
        self._add_executive_summary(results)
        self._add_table_of_contents()

        for result in results:
            self._add_website_assessment(result, screenshots_dir)

        self._add_overall_recommendations(results)
        self._add_manual_review_checklist(results)

        # Save document
        self.doc.save(str(self.config.output_path))
        return self.config.output_path

    def _add_title_page(self):
        """Add title page."""
        # Add logo placeholder
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        run = title_para.add_run(self.config.name)
        run.font.name = "Calibri"
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Website Optimisation Assessment Report")
        subtitle_run.font.name = "Calibri"
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Date
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        date_run.font.name = "Calibri"
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Page break
        self.doc.add_page_break()

    def _add_executive_summary(self, results: list[AssessmentResult]):
        """Add executive summary section."""
        self.doc.add_heading("Executive Summary", level=1)
        
        summary = self.doc.add_paragraph()
        summary.add_run(
            f"This report presents a comprehensive assessment of {len(results)} website(s), "
            f"evaluating technical performance, user experience, content quality, SEO, "
            f"and analytics implementation. The assessment follows industry best practices "
            f"and provides actionable recommendations for improvement."
        )
        
        # Summary statistics
        self.doc.add_heading("Assessment Overview", level=2)
        
        total_findings = 0
        severity_counts = {Severity.CRITICAL: 0, Severity.HIGH: 0, Severity.MEDIUM: 0, Severity.LOW: 0}
        
        for result in results:
            for section in result.sections.values():
                for finding in section.findings:
                    total_findings += 1
                    severity_counts[finding.severity] = severity_counts.get(finding.severity, 0) + 1
        
        # Create summary table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        
        severity_labels = {
            Severity.CRITICAL: "Critical",
            Severity.HIGH: "High",
            Severity.MEDIUM: "Medium",
            Severity.LOW: "Low",
        }
        
        for i, (severity, count) in enumerate(severity_counts.items()):
            row = table.rows[i]
            row.cells[0].text = severity_labels[severity]
            row.cells[1].text = str(count)
        
        self.doc.add_paragraph()

    def _add_table_of_contents(self):
        """Add table of contents placeholder."""
        self.doc.add_heading("Table of Contents", level=1)
        self.doc.add_paragraph(
            "Note: In Microsoft Word, update the table of contents by right-clicking "
            "and selecting 'Update Field'."
        )
        self.doc.add_page_break()

    def _add_website_assessment(self, result: AssessmentResult, screenshots_dir: Path):
        """Add assessment section for a single website."""
        # Website heading
        self.doc.add_heading(result.website_name, level=1)
        
        # URL
        url_para = self.doc.add_paragraph()
        url_run = url_para.add_run("URL: ")
        url_run.bold = True
        url_para.add_run(result.website_url)
        
        # Overall score
        if result.overall_score > 0:
            score_para = self.doc.add_paragraph()
            score_run = score_para.add_run(f"Overall Score: {result.overall_score:.1f}/100")
            score_run.bold = True
            score_run.font.color.rgb = self._get_score_color(result.overall_score)
        
        self.doc.add_paragraph()
        
        # Add sections
        for section_name, section_result in result.sections.items():
            self._add_section(section_name, section_result, screenshots_dir)

    def _add_section(self, section_name: str, section: SectionResult, screenshots_dir: Path):
        """Add a section to the report."""
        self.doc.add_heading(section_name, level=2)
        
        # Findings
        if section.findings:
            self.doc.add_heading("Findings", level=3)
            
            for finding in section.findings:
                self._add_finding(finding)
        
        # Screenshots
        if section.screenshots:
            self.doc.add_heading("Screenshots", level=3)

            for screenshot_path in section.screenshots:
                if screenshot_path.exists():
                    self.doc.add_picture(str(screenshot_path), width=Inches(6.0))
                    self.doc.add_paragraph(screenshot_path.name, style="Caption")
        
        # Recommendations
        if section.recommendations:
            self.doc.add_heading("Recommendations", level=3)
            
            for rec in section.recommendations:
                self._add_recommendation(rec)
        
        self.doc.add_paragraph()

    def _add_finding(self, finding: Finding):
        """Add a finding to the report."""
        # Status indicator
        status_symbols = {
            Status.PASS: "✓",
            Status.WARNING: "⚠",
            Status.FAIL: "✗",
            Status.NOT_CHECKED: "?",
        }
        
        severity_colors = {
            Severity.CRITICAL: RGBColor(139, 0, 0),
            Severity.HIGH: RGBColor(200, 0, 0),
            Severity.MEDIUM: RGBColor(255, 140, 0),
            Severity.LOW: RGBColor(0, 100, 0),
        }
        
        para = self.doc.add_paragraph()
        
        # Status symbol
        status_run = para.add_run(f"{status_symbols.get(finding.status, '?')} ")
        
        # Description
        desc_run = para.add_run(finding.description)
        desc_run.bold = True
        
        # Severity
        sev_run = para.add_run(f" [{finding.severity.value.upper()}]")
        sev_run.font.color.rgb = severity_colors.get(finding.severity, RGBColor(100, 100, 100))
        sev_run.font.size = Pt(9)
        
        # Evidence
        if finding.evidence:
            evidence_para = self.doc.add_paragraph()
            evidence_run = evidence_para.add_run(f"Evidence: {finding.evidence}")
            evidence_run.font.size = Pt(10)
            evidence_run.font.color.rgb = RGBColor(100, 100, 100)
            evidence_para.paragraph_format.left_indent = Inches(0.5)

    def _add_recommendation(self, rec: Recommendation):
        """Add a recommendation to the report."""
        priority_symbols = {
            "high": "🔴",
            "medium": "🟡",
            "low": "🟢",
        }
        
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        
        # Priority indicator
        priority_run = para.add_run(f"{priority_symbols.get(rec.priority, '•')} ")
        
        # Description
        desc_run = para.add_run(rec.description)
        
        # Effort
        effort_run = para.add_run(f" (Effort: {rec.effort})")
        effort_run.font.size = Pt(10)
        effort_run.font.color.rgb = RGBColor(100, 100, 100)

    def _add_overall_recommendations(self, results: list[AssessmentResult]):
        """Add overall recommendations section."""
        self.doc.add_heading("Overall Recommendations", level=1)

        self.doc.add_paragraph(
            "Based on the comprehensive assessment, the following priority actions are recommended:"
        )

        # Collect high-priority recommendations
        high_priority = []
        for result in results:
            for section in result.sections.values():
                for rec in section.recommendations:
                    if rec.priority == "high":
                        high_priority.append((result.website_name, rec))

        if high_priority:
            self.doc.add_heading("High Priority Actions", level=2)
            for website_name, rec in high_priority[:10]:  # Top 10
                para = self.doc.add_paragraph(style="List Bullet")
                para.add_run(f"[{website_name}] {rec.description}")

        # Next steps
        self.doc.add_heading("Next Steps", level=2)
        next_steps = [
            "Review findings with stakeholders",
            "Prioritize recommendations based on business impact",
            "Create implementation timeline",
            "Assign ownership for each action item",
            "Schedule follow-up assessment in 3-6 months",
        ]

        for step in next_steps:
            self.doc.add_paragraph(step, style="List Bullet")

    def _add_manual_review_checklist(self, results: list[AssessmentResult]):
        """Add manual review checklist for items requiring human assessment."""
        self.doc.add_page_break()
        self.doc.add_heading("Manual Review Checklist", level=1)
        
        self.doc.add_paragraph(
            "The following items require manual review as they cannot be fully automated. "
            "Use the provided credentials and links to complete the assessment."
        )
        
        # Credentials section
        if self.config.credentials:
            self.doc.add_heading("WordPress Admin Access", level=2)
            self.doc.add_paragraph(
                f"URL: {self.config.information_urls[0].url if self.config.information_urls else 'N/A'}",
                style="Intense Quote"
            )
            self.doc.add_paragraph(
                f"Username: {self.config.credentials.username}\n"
                f"Password: {self.config.credentials.password}",
                style="No Spacing"
            )
        
        # Analytics access
        analytics_urls = [u for u in self.config.information_urls if "analytics" in u.url.lower()]
        if analytics_urls:
            self.doc.add_heading("Google Analytics", level=2)
            for url in analytics_urls:
                self.doc.add_paragraph(f"URL: {url.url}", style="Intense Quote")
        
        # Checklist items
        self.doc.add_heading("Items Requiring Manual Review", level=2)
        
        checklist_items = [
            ("Visual Design & Layout", [
                "Check for layout breakages on mobile devices",
                "Verify button sizes are touch-friendly (min 44x44px)",
                "Assess color contrast for WCAG AA compliance",
                "Review overall visual design and branding consistency",
            ]),
            ("Content Quality", [
                "Check for outdated staff/team member information",
                "Review product/service descriptions for accuracy",
                "Identify old blog posts that need updating or archiving",
                "Assess messaging consistency across all pages",
            ]),
            ("WordPress Security", [
                "Verify no default admin usernames (admin, administrator)",
                "Check for weak passwords or shared accounts",
                "Review user roles and permissions",
                "Check for publicly visible error messages",
            ]),
            ("Technical SEO", [
                "Review full site for duplicate content issues",
                "Check redirect chains using browser dev tools",
                "Verify canonical tags on all important pages",
                "Review internal linking structure",
            ]),
            ("Analytics & Tracking", [
                "Log into Google Analytics to verify goals/events setup",
                "Check for tracking gaps in conversion funnels",
                "Verify Google Tag Manager container is firing correctly",
                "Review cookie consent integration with analytics",
            ]),
            ("Navigation & UX", [
                "Test complete user journeys for dead ends",
                "Identify confusing menu items or navigation paths",
                "Test keyboard navigation throughout the site",
                "Verify search functionality (if available)",
            ]),
        ]
        
        for category, items in checklist_items:
            self.doc.add_heading(category, level=3)
            for item in items:
                para = self.doc.add_paragraph(style="List Bullet")
                para.add_run(f"☐ {item}")  # Empty checkbox
        
        # Notes section
        self.doc.add_heading("Manual Review Notes", level=2)
        self.doc.add_paragraph(
            "Use the space below to record findings from manual review:",
            style="No Spacing"
        )
        
        # Add blank lines for notes
        for _ in range(10):
            self.doc.add_paragraph("_", style="No Spacing")

    def _get_score_color(self, score: float) -> RGBColor:
        """Get color based on score."""
        if score >= 80:
            return RGBColor(0, 100, 0)  # Green
        elif score >= 60:
            return RGBColor(255, 140, 0)  # Orange
        else:
            return RGBColor(200, 0, 0)  # Red
