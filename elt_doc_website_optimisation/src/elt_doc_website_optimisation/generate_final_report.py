"""Generate enhanced client-ready report - with screenshots and full appendices."""

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_findings(raw_doc_path: Path) -> dict:
    """Extract actual findings from Python assessment."""
    doc = Document(raw_doc_path)
    
    findings = {
        'cnltd': {
            'technical': [], 'ux': [], 'content': [], 'seo': [],
            'wordpress': [], 'analytics': [], 'visual': [], 'score': 0,
        },
        'insurance': {
            'technical': [], 'ux': [], 'content': [], 'seo': [],
            'wordpress': [], 'analytics': [], 'visual': [], 'score': 0,
        },
    }
    
    current_site = None
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        if 'cnltd.co.uk' in text.lower():
            current_site = 'cnltd'
        elif 'insuranceit.co.uk' in text.lower():
            current_site = 'insurance'
        
        if 'Technical Review' in text and 'Heading' in str(para.style):
            current_section = 'technical'
        elif 'UX' in text and 'Heading' in str(para.style):
            current_section = 'ux'
        elif 'Content' in text and 'Heading' in str(para.style):
            current_section = 'content'
        elif 'SEO' in text and 'Heading' in str(para.style):
            current_section = 'seo'
        elif 'Plugin' in text and 'Heading' in str(para.style):
            current_section = 'wordpress'
        elif 'Analytics' in text and 'Heading' in str(para.style):
            current_section = 'analytics'
        elif 'Visual' in text and 'Heading' in str(para.style):
            current_section = 'visual'
        
        if current_site and current_section:
            if any(text.startswith(s) for s in ['✓', '⚠', '✗', '?']):
                finding = {'text': text, 'evidence': '', 'status': 'pass' if text.startswith('✓') else 'warning' if text.startswith('⚠') else 'fail'}
                findings[current_site][current_section].append(finding)
            elif text.startswith('Evidence:'):
                if findings[current_site][current_section]:
                    findings[current_site][current_section][-1]['evidence'] = text.replace('Evidence:', '').strip()
            elif 'Overall Score' in text and current_site:
                try:
                    score = float(text.split('Overall Score:')[1].split('/')[0].strip())
                    findings[current_site]['score'] = score
                except:
                    pass
    
    return findings


def create_enhanced_report(output_path: Path, raw_doc_path: Path, screenshots_dir: Path):
    """Create the final client-ready enhanced report with screenshots and appendices."""
    
    findings = extract_findings(raw_doc_path)
    doc = Document()
    
    # === TITLE PAGE ===
    for _ in range(5):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Website Optimisation\nAssessment Report")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    for _ in range(3):
        doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Commercial Networks LTD & Insurance IT")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(8):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # === TABLE OF CONTENTS ===
    doc.add_heading("Table of Contents", level=1)
    toc = """
1. Executive Summary
2. Assessment Overview
3. Technical Review
   3.1 Performance
   3.2 Security
   3.3 Hosting & Infrastructure
4. UX & Navigation
   4.1 Navigation Flow
   4.2 Accessibility
   4.3 Mobile Friendliness
   4.4 Desktop Screenshots
   4.5 Mobile Screenshots
5. Content & Messaging
   5.1 Tone and Clarity
   5.2 Calls to Action
   5.3 Outdated Content
6. SEO Review
   6.1 On-page SEO
   6.2 Technical SEO
   6.3 Content SEO
7. Plugin & Theme Audit
   7.1 Plugin Review
   7.2 Theme Review
   7.3 Recommendations
8. Analytics & Tracking
9. Priority Action Plan
10. Manual Review Checklist
Appendices
   Appendix A: Technical Review Findings
   Appendix B: UX & Navigation Findings
   Appendix C: Content & Messaging Findings
   Appendix D: SEO Review Findings
   Appendix E: Plugin & Theme Audit Findings
   Appendix F: Analytics & Tracking Findings
   Appendix G: Visual Analysis Findings
"""
    doc.add_paragraph(toc)
    doc.add_page_break()
    
    # === 1. EXECUTIVE SUMMARY ===
    doc.add_heading("1. Executive Summary", level=1)
    
    cnltd_score = findings['cnltd']['score']
    insurance_score = findings['insurance']['score']
    
    exec_summary = f"""This report presents a comprehensive website optimisation assessment for Commercial Networks LTD (cnltd.co.uk) and Insurance IT (insuranceit.co.uk). The assessment evaluated both websites against the full specification requirements covering six key dimensions: Technical Performance, User Experience & Navigation, Content & Messaging, Search Engine Optimisation (SEO), WordPress Plugin & Theme Configuration, and Analytics & Tracking implementation.

Overall, both websites demonstrate solid technical foundations with scores of {cnltd_score}/100 and {insurance_score}/100 respectively. The assessment identified several strengths including responsive mobile design, proper HTTPS configuration, current WordPress version (6.9.4), and comprehensive plugin installations (31 plugins on cnltd.co.uk, 10 on insuranceit.co.uk).

Critical findings requiring immediate attention include: missing security headers on both websites (HSTS, X-Content-Type-Options, X-Frame-Options, Content-Security-Policy), images without alt text affecting accessibility compliance (1 on cnltd.co.uk, 9 on insuranceit.co.uk), cookie consent banners without clear reject options (GDPR concern), and no clear calls-to-action detected on insuranceit.co.uk.

The detailed findings in this report include evidence from automated testing, screenshots demonstrating mobile and desktop rendering, and a complete appendix with all assessment data. Priority actions are categorised by effort and impact to support implementation planning."""
    
    for paragraph in exec_summary.split("\n\n"):
        doc.add_paragraph(paragraph)
    
    # === 2. ASSESSMENT OVERVIEW ===
    doc.add_heading("2. Assessment Overview", level=1)
    
    doc.add_heading("2.1 Websites Assessed", level=2)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    
    headers = table.rows[0].cells
    headers[0].text = "Website"
    headers[1].text = "URL"
    headers[2].text = "Score"
    headers[3].text = "Status"
    for cell in headers:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    
    data = [
        ("Commercial Networks LTD", "cnltd.co.uk", f"{cnltd_score}/100", "Good" if cnltd_score >= 75 else "Fair"),
        ("Insurance IT", "insuranceit.co.uk", f"{insurance_score}/100", "Good" if insurance_score >= 75 else "Fair"),
    ]
    for i, (name, url, score, status) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = url
        row.cells[2].text = score
        row.cells[3].text = status
    
    # === 3. TECHNICAL REVIEW ===
    doc.add_heading("3. Technical Review", level=1)
    doc.add_paragraph("This section evaluates the technical performance, security configuration, and hosting infrastructure of both websites. The assessment includes server response times, SSL/TLS configuration, security headers, WordPress version verification, and plugin analysis.")
    
    doc.add_heading("3.1 Performance", level=2)
    doc.add_paragraph("Performance analysis measures server response times and page load characteristics. Target response time is under 1 second for optimal user experience and SEO ranking.")
    
    doc.add_heading("3.2 Security", level=2)
    doc.add_paragraph("Security assessment verifies WordPress core version, plugin versions, SSL certificate configuration, and security header implementation. Missing security headers expose websites to common vulnerabilities including clickjacking, MIME-type sniffing, and protocol downgrade attacks.")
    
    doc.add_heading("3.3 Hosting & Infrastructure", level=2)
    doc.add_paragraph("Hosting evaluation examines server response times and configuration settings. PHP version verification requires WordPress admin access.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['technical']:
            doc.add_paragraph("No technical findings recorded.")
        for finding in site_data['technical']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # === 4. UX & NAVIGATION ===
    doc.add_heading("4. UX & Navigation", level=1)
    doc.add_paragraph("User Experience and Navigation assessment evaluates website usability, accessibility compliance, and mobile responsiveness. This section examines navigation structure, image accessibility, heading hierarchy, and mobile viewport configuration.")
    
    doc.add_heading("4.1 Navigation Flow", level=2)
    doc.add_paragraph("Navigation analysis verifies semantic HTML5 structure, link organization, and user journey flow. Well-structured navigation improves both user experience and SEO performance.")
    
    doc.add_heading("4.2 Accessibility", level=2)
    doc.add_paragraph("Accessibility assessment checks WCAG compliance including image alt text, heading structure, color contrast, and keyboard navigation support. Accessible websites serve all users and meet legal requirements.")
    
    doc.add_heading("4.3 Mobile Friendliness", level=2)
    doc.add_paragraph("Mobile assessment verifies responsive design implementation, viewport configuration, and touch-friendly interface elements. Mobile-first indexing makes mobile optimization critical for SEO.")
    
    doc.add_heading("4.4 Desktop Screenshots (1920×1080)", level=2)
    doc.add_paragraph("Desktop screenshots demonstrate website rendering at standard desktop resolution (1920×1080 pixels). Review layout, navigation visibility, and content alignment.")
    
    doc.add_heading("4.5 Mobile Screenshots (375×667)", level=2)
    doc.add_paragraph("Mobile screenshots demonstrate responsive design at smartphone resolution (375×667 pixels, iPhone SE size). Review mobile menu activation, content readability, and touch target sizing.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['ux']:
            doc.add_paragraph("No UX findings recorded.")
        for finding in site_data['ux']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # 4.4 Desktop Screenshots
    doc.add_heading("4.4 Desktop Screenshots (1920×1080)", level=2)
    
    doc.add_heading("Commercial Networks LTD", level=3)
    desktop_cnltd = screenshots_dir / "Website to assess_desktop.png"
    if desktop_cnltd.exists():
        try:
            doc.add_picture(str(desktop_cnltd), width=Inches(6.5))
            doc.add_paragraph("Figure 1: Commercial Networks LTD - Desktop View", style="Caption")
        except:
            doc.add_paragraph(f"Screenshot available at: {desktop_cnltd}")
    else:
        doc.add_paragraph(f"Screenshot location: {desktop_cnltd}")
    
    doc.add_heading("Insurance IT", level=3)
    desktop_ins = screenshots_dir / "Website to assess_desktop.png"
    if desktop_ins.exists():
        try:
            doc.add_picture(str(desktop_ins), width=Inches(6.5))
            doc.add_paragraph("Figure 2: Insurance IT - Desktop View", style="Caption")
        except:
            doc.add_paragraph(f"Screenshot available at: {desktop_ins}")
    else:
        doc.add_paragraph(f"Screenshot location: {desktop_ins}")
    
    # 4.5 Mobile Screenshots
    doc.add_heading("4.5 Mobile Screenshots (375×667)", level=2)
    
    doc.add_heading("Commercial Networks LTD", level=3)
    mobile_cnltd = screenshots_dir / "Website to assess_mobile.png"
    if mobile_cnltd.exists():
        try:
            doc.add_picture(str(mobile_cnltd), width=Inches(3.0))
            doc.add_paragraph("Figure 3: Commercial Networks LTD - Mobile View", style="Caption")
        except:
            doc.add_paragraph(f"Screenshot available at: {mobile_cnltd}")
    else:
        doc.add_paragraph(f"Screenshot location: {mobile_cnltd}")
    
    doc.add_heading("Insurance IT", level=3)
    mobile_ins = screenshots_dir / "Website to assess_mobile.png"
    if mobile_ins.exists():
        try:
            doc.add_picture(str(mobile_ins), width=Inches(3.0))
            doc.add_paragraph("Figure 4: Insurance IT - Mobile View", style="Caption")
        except:
            doc.add_paragraph(f"Screenshot available at: {mobile_ins}")
    else:
        doc.add_paragraph(f"Screenshot location: {mobile_ins}")
    
    # === 5. CONTENT & MESSAGING ===
    doc.add_heading("5. Content & Messaging", level=1)
    doc.add_paragraph("Content assessment evaluates messaging clarity, call-to-action effectiveness, and content freshness. Quality content drives user engagement, conversions, and SEO performance.")
    
    doc.add_heading("5.1 Tone and Clarity", level=2)
    doc.add_paragraph("Content analysis measures word count, value proposition clarity, and text density. Optimal content balances information with readability.")
    
    doc.add_heading("5.2 Calls to Action", level=2)
    doc.add_paragraph("Call-to-action assessment identifies conversion opportunities and user guidance. Every page should have at least one clear, relevant CTA.")
    
    doc.add_heading("5.3 Outdated Content", level=2)
    doc.add_paragraph("Content freshness analysis detects outdated dates, staff information, and product/service descriptions. Fresh content signals active business operations to users and search engines.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['content']:
            doc.add_paragraph("No content findings recorded.")
        for finding in site_data['content']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # === 6. SEO REVIEW ===
    doc.add_heading("6. SEO Review", level=1)
    doc.add_paragraph("Search Engine Optimization assessment evaluates on-page elements, technical SEO implementation, and content optimization. Strong SEO foundations improve search visibility and organic traffic.")
    
    doc.add_heading("6.1 On-page SEO", level=2)
    doc.add_paragraph("On-page SEO analysis examines title tags, meta descriptions, heading structure, image alt text, and URL structure. These elements directly impact search rankings and click-through rates.")
    
    doc.add_heading("6.2 Technical SEO", level=2)
    doc.add_paragraph("Technical SEO assessment verifies link integrity, redirect chains, sitemap availability, and robots.txt configuration. Technical issues can prevent search engines from properly crawling and indexing content.")
    
    doc.add_heading("6.3 Content SEO", level=2)
    doc.add_paragraph("Content SEO evaluation identifies thin content, duplicate content issues, and content expansion opportunities. Quality, unique content is essential for search visibility.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['seo']:
            doc.add_paragraph("No SEO findings recorded.")
        for finding in site_data['seo']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # === 7. PLUGIN & THEME AUDIT ===
    doc.add_heading("7. Plugin & Theme Audit", level=1)
    doc.add_paragraph("WordPress plugin and theme assessment evaluates installed extensions, version currency, and security posture. Plugins extend functionality but introduce security and maintenance considerations.")
    
    doc.add_heading("7.1 Plugin Review", level=2)
    doc.add_paragraph("Plugin analysis identifies installed extensions, version information, and potential security concerns. Outdated or vulnerable plugins are a primary attack vector for WordPress sites.")
    
    doc.add_heading("7.2 Theme Review", level=2)
    doc.add_paragraph("Theme assessment verifies theme identity, support status, and customization documentation. Themes control visual presentation and impact performance and security.")
    
    doc.add_heading("7.3 Recommendations", level=2)
    doc.add_paragraph("Plugin and theme recommendations prioritise security, performance, and maintainability. Regular updates and consolidation reduce attack surface and maintenance burden.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['wordpress']:
            doc.add_paragraph("No WordPress findings recorded.")
        for finding in site_data['wordpress']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # === 8. ANALYTICS & TRACKING ===
    doc.add_heading("8. Analytics & Tracking", level=1)
    doc.add_paragraph("Analytics and tracking assessment verifies Google Analytics implementation, Tag Manager usage, and cookie consent compliance. Proper tracking enables data-driven decision making while respecting user privacy.")
    
    for site_name, site_data in [('Commercial Networks LTD', findings['cnltd']), ('Insurance IT', findings['insurance'])]:
        doc.add_heading(f"Findings for {site_name}", level=3)
        if not site_data['analytics']:
            doc.add_paragraph("No analytics findings recorded.")
        for finding in site_data['analytics']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # === 9. PRIORITY ACTION PLAN ===
    doc.add_heading("9. Priority Action Plan", level=1)
    
    doc.add_heading("9.1 Critical Actions (Immediate)", level=2)
    critical = [
        "Add security headers (HSTS, X-Content-Type-Options, X-Frame-Options, CSP) - 1-2 hours",
        "Add alt text to all images (1 on cnltd, 9 on insuranceit) - 2-3 hours",
        "Ensure cookie banner has clear reject option (GDPR) - 1 hour",
        "Add clear calls-to-action to insuranceit.co.uk - 2 hours",
    ]
    for item in critical:
        doc.add_paragraph(f"✗ {item}", style="List Bullet")
    
    doc.add_heading("9.2 High Priority (This Week)", level=2)
    high = [
        "Review robots.txt configuration (Disallow: / detected) - 30 minutes",
        "Fix duplicate page titles - 1 hour",
        "Update WordPress plugins without version info - 1 hour",
        "Verify WordPress admin security (no default usernames) - 30 minutes",
    ]
    for item in high:
        doc.add_paragraph(f"⚠ {item}", style="List Bullet")
    
    doc.add_heading("9.3 Medium Priority (This Month)", level=2)
    medium = [
        "Review and consolidate 31 plugins on cnltd.co.uk - 1-2 days",
        "Update outdated content (dates from 2021) - 2-3 hours",
        "Set up GA4 conversion tracking - 2 hours",
        "Add sitemap reference to robots.txt - 30 minutes",
    ]
    for item in medium:
        doc.add_paragraph(f"○ {item}", style="List Bullet")
    
    # === 10. MANUAL REVIEW CHECKLIST ===
    doc.add_page_break()
    doc.add_heading("10. Items Requiring Manual Review", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Note: ").bold = True
    p.add_run("The following items require human review as they cannot be fully automated. WordPress admin credentials are provided separately via secure channel.")
    
    checklist = [
        ("Visual Design", [
            "Check for layout breakages on mobile devices",
            "Verify button sizes are touch-friendly (min 44x44px)",
            "Assess color contrast for WCAG AA compliance",
        ]),
        ("Content Quality", [
            "Check for outdated staff/team member information",
            "Review product/service descriptions for accuracy",
            "Identify old blog posts that need updating or archiving",
        ]),
        ("WordPress Security", [
            "Verify no default admin usernames (admin, administrator)",
            "Check for weak passwords or shared accounts",
            "Review user roles and permissions",
            "Check PHP version via Site Health dashboard",
        ]),
        ("Analytics", [
            "Log into Google Analytics to verify goals/events setup",
            "Check for tracking gaps in conversion funnels",
            "Verify cookie consent integration with analytics",
        ]),
    ]
    
    for category, items in checklist:
        doc.add_heading(category, level=2)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"☐ {item}")
    
    # === APPENDICES ===
    doc.add_page_break()
    doc.add_heading("Appendices", level=1)
    doc.add_paragraph("The appendices contain complete findings from all automated assessments. Each appendix corresponds to a specification section and includes all evidence collected during testing.")
    
    # Appendix A: Technical Review
    doc.add_heading("Appendix A: Technical Review Findings", level=2)
    doc.add_paragraph("Complete technical assessment findings including performance metrics, security configuration, and hosting infrastructure analysis.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['technical']:
            doc.add_paragraph("No technical findings recorded.")
        for finding in site_data['technical']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix B: UX & Navigation
    doc.add_heading("Appendix B: UX & Navigation Findings", level=2)
    doc.add_paragraph("Complete UX assessment findings including navigation structure, accessibility compliance, and mobile responsiveness analysis.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['ux']:
            doc.add_paragraph("No UX findings recorded.")
        for finding in site_data['ux']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix C: Content & Messaging
    doc.add_heading("Appendix C: Content & Messaging Findings", level=2)
    doc.add_paragraph("Complete content assessment findings including word count analysis, CTA detection, and content freshness evaluation.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['content']:
            doc.add_paragraph("No content findings recorded.")
        for finding in site_data['content']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix D: SEO Review
    doc.add_heading("Appendix D: SEO Review Findings", level=2)
    doc.add_paragraph("Complete SEO assessment findings including on-page elements, technical SEO, and content optimization analysis.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['seo']:
            doc.add_paragraph("No SEO findings recorded.")
        for finding in site_data['seo']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix E: Plugin & Theme Audit
    doc.add_heading("Appendix E: Plugin & Theme Audit Findings", level=2)
    doc.add_paragraph("Complete WordPress plugin and theme assessment findings including version information and security analysis.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['wordpress']:
            doc.add_paragraph("No WordPress findings recorded.")
        for finding in site_data['wordpress']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix F: Analytics & Tracking
    doc.add_heading("Appendix F: Analytics & Tracking Findings", level=2)
    doc.add_paragraph("Complete analytics and tracking assessment findings including GA4, GTM, and cookie compliance verification.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['analytics']:
            doc.add_paragraph("No analytics findings recorded.")
        for finding in site_data['analytics']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Appendix G: Visual Analysis
    doc.add_heading("Appendix G: Visual Analysis Findings", level=2)
    doc.add_paragraph("Complete visual analysis findings including color contrast, font sizing, and link styling assessment.")
    for site_name, site_data in [('Commercial Networks LTD (cnltd.co.uk)', findings['cnltd']), ('Insurance IT (insuranceit.co.uk)', findings['insurance'])]:
        doc.add_heading(site_name, level=3)
        if not site_data['visual']:
            doc.add_paragraph("No visual findings recorded.")
        for finding in site_data['visual']:
            p = doc.add_paragraph(style="List Bullet")
            text = finding['text'].lstrip('✓⚠✗? ').strip()
            symbol = "✓" if finding['status'] == 'pass' else "⚠" if finding['status'] == 'warning' else "✗"
            p.add_run(f"{symbol} {text}").bold = True
            if finding['evidence']:
                p.add_run(f"\n    Evidence: {finding['evidence']}")
    
    # Save
    doc.save(str(output_path))
    print(f"✓ Enhanced client-ready report saved to: {output_path}")


if __name__ == "__main__":
    output = Path("/Users/rpatel/Downloads/website_optimisation_FINAL_REPORT.docx")
    raw = Path("/Users/rpatel/Downloads/website_optimisation_assessment.docx")
    screenshots = Path.home() / "Downloads" / "website_assessment_screenshots"
    create_enhanced_report(output, raw, screenshots)
